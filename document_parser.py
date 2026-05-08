"""储能配置AGENT - 文档解析模块
支持读取PDF、Word、图片等各类文件，提取电费数据。
"""

import logging
from pathlib import Path
from typing import Optional

import pandas as pd

logger = logging.getLogger(__name__)


class DocumentParser:
    """多格式文档解析器"""

    SUPPORTED_EXTENSIONS = {
        ".pdf", ".doc", ".docx", ".xls", ".xlsx",
        ".csv", ".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif",
        ".txt",
    }

    def __init__(self, ocr_language: str = "ch_sim+en", tesseract_lang: str = "chi_sim+eng"):
        self.ocr_language = ocr_language
        self.tesseract_lang = tesseract_lang
        self._easyocr_reader = None

    # ------------------------------------------------------------------
    # 公共入口
    # ------------------------------------------------------------------
    def parse(self, file_path: str | Path) -> dict:
        """解析单个文件，返回统一结构。

        Returns:
            {
                "file": str,
                "type": str,           # pdf / word / image / excel / csv / text
                "text": str,           # 提取到的全文文本
                "tables": list[pd.DataFrame],  # 提取到的表格
                "pages": list[dict],   # 逐页内容 (PDF/Word)
            }
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")
        if path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件类型: {path.suffix}")

        ext = path.suffix.lower()
        if ext == ".pdf":
            return self._parse_pdf(path)
        elif ext in (".doc", ".docx"):
            return self._parse_word(path)
        elif ext in (".xls", ".xlsx"):
            return self._parse_excel(path)
        elif ext == ".csv":
            return self._parse_csv(path)
        elif ext == ".txt":
            return self._parse_text(path)
        elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif"):
            return self._parse_image(path)
        else:
            raise ValueError(f"未处理的文件类型: {ext}")

    def parse_batch(self, file_paths: list[str | Path]) -> list[dict]:
        """批量解析多个文件。"""
        results = []
        for fp in file_paths:
            try:
                result = self.parse(fp)
                results.append(result)
            except Exception as e:
                logger.warning("解析文件 %s 失败: %s", fp, e)
                results.append({"file": str(fp), "error": str(e)})
        return results

    # ------------------------------------------------------------------
    # PDF 解析
    # ------------------------------------------------------------------
    def _parse_pdf(self, path: Path) -> dict:
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("请安装 pdfplumber: pip install pdfplumber")

        pages = []
        all_text_parts = []
        all_tables = []

        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                page_tables = page.extract_tables()

                # 将表格转为 DataFrame
                dfs = []
                for tbl in page_tables:
                    if tbl and len(tbl) > 1:
                        df = pd.DataFrame(tbl[1:], columns=tbl[0])
                        dfs.append(df)
                        all_tables.append(df)

                pages.append({
                    "page": i + 1,
                    "text": page_text,
                    "tables": dfs,
                })
                all_text_parts.append(page_text)

        # 如果 PDF 无可提取文本，尝试 OCR
        full_text = "\n".join(all_text_parts).strip()
        if not full_text:
            logger.info("PDF无可提取文本，尝试OCR...")
            full_text, ocr_tables = self._ocr_pdf(path)
            all_tables.extend(ocr_tables)

        return {
            "file": str(path),
            "type": "pdf",
            "text": full_text,
            "tables": all_tables,
            "pages": pages,
        }

    def _ocr_pdf(self, path: Path) -> tuple[str, list[pd.DataFrame]]:
        """将PDF转图片后进行OCR识别。"""
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(str(path))
        except ImportError:
            logger.warning("pdf2image未安装，无法对PDF进行OCR。请安装: pip install pdf2image")
            return "", []

        full_text_parts = []
        for img in images:
            text = self._ocr_image_obj(img)
            full_text_parts.append(text)
        return "\n".join(full_text_parts), []

    # ------------------------------------------------------------------
    # Word 解析
    # ------------------------------------------------------------------
    def _parse_word(self, path: Path) -> dict:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = Document(str(path))
        pages = []
        all_text_parts = []
        all_tables = []

        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                all_text_parts.append(para.text)

        # 提取表格
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows and len(rows) > 1:
                df = pd.DataFrame(rows[1:], columns=rows[0])
                all_tables.append(df)

        pages.append({
            "page": 1,
            "text": "\n".join(all_text_parts),
            "tables": all_tables,
        })

        return {
            "file": str(path),
            "type": "word",
            "text": "\n".join(all_text_parts),
            "tables": all_tables,
            "pages": pages,
        }

    # ------------------------------------------------------------------
    # Excel / CSV 解析
    # ------------------------------------------------------------------
    def _parse_excel(self, path: Path) -> dict:
        all_tables = []
        all_text_parts = []

        xls = pd.ExcelFile(str(path))
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            all_tables.append(df)
            all_text_parts.append(f"[Sheet: {sheet_name}]\n{df.to_string()}")

        return {
            "file": str(path),
            "type": "excel",
            "text": "\n".join(all_text_parts),
            "tables": all_tables,
            "pages": [],
        }

    def _parse_csv(self, path: Path) -> dict:
        df = pd.read_csv(str(path))
        return {
            "file": str(path),
            "type": "csv",
            "text": df.to_string(),
            "tables": [df],
            "pages": [],
        }

    # ------------------------------------------------------------------
    # 纯文本解析
    # ------------------------------------------------------------------
    def _parse_text(self, path: Path) -> dict:
        text = path.read_text(encoding="utf-8", errors="replace")
        return {
            "file": str(path),
            "type": "text",
            "text": text,
            "tables": [],
            "pages": [{"page": 1, "text": text, "tables": []}],
        }

    # ------------------------------------------------------------------
    # 图片 OCR 解析
    # ------------------------------------------------------------------
    def _parse_image(self, path: Path) -> dict:
        from PIL import Image
        img = Image.open(str(path))
        text = self._ocr_image_obj(img)
        tables = self._try_parse_text_table(text)

        return {
            "file": str(path),
            "type": "image",
            "text": text,
            "tables": tables,
            "pages": [{"page": 1, "text": text, "tables": tables}],
        }

    def _ocr_image_obj(self, img) -> str:
        """对 PIL Image 对象执行 OCR，优先用 EasyOCR，回退到 Tesseract。"""
        # 尝试 EasyOCR
        text = self._ocr_with_easyocr(img)
        if text.strip():
            return text

        # 回退到 Tesseract
        text = self._ocr_with_tesseract(img)
        return text

    def _ocr_with_easyocr(self, img) -> str:
        try:
            if self._easyocr_reader is None:
                import easyocr
                langs = self.ocr_language.split("+")
                self._easyocr_reader = easyocr.Reader(langs, gpu=False)

            import numpy as np
            img_array = np.array(img)
            results = self._easyocr_reader.readtext(img_array, detail=0)
            return "\n".join(results)
        except ImportError:
            logger.debug("easyocr未安装，跳过")
            return ""
        except Exception as e:
            logger.warning("EasyOCR识别失败: %s", e)
            return ""

    def _ocr_with_tesseract(self, img) -> str:
        try:
            import pytesseract
            text = pytesseract.image_to_string(img, lang=self.tesseract_lang)
            return text
        except ImportError:
            logger.warning("pytesseract未安装，跳过OCR")
            return ""
        except Exception as e:
            logger.warning("Tesseract识别失败: %s", e)
            return ""

    # ------------------------------------------------------------------
    # 辅助：从纯文本尝试解析表格
    # ------------------------------------------------------------------
    @staticmethod
    def _try_parse_text_table(text: str) -> list[pd.DataFrame]:
        """尝试从OCR文本中识别表格结构（简单的制表符/多空格分隔）。"""
        import re
        lines = text.strip().split("\n")
        if len(lines) < 2:
            return []

        # 检测分隔符
        table_lines = []
        for line in lines:
            # 用两个以上空格或制表符分割
            parts = re.split(r"\t|\s{2,}", line.strip())
            if len(parts) >= 2:
                table_lines.append(parts)

        if len(table_lines) < 2:
            return []

        # 统一列数
        max_cols = max(len(row) for row in table_lines)
        normalized = [row + [""] * (max_cols - len(row)) for row in table_lines]

        try:
            df = pd.DataFrame(normalized[1:], columns=normalized[0])
            return [df]
        except Exception:
            return []
