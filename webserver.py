"""FastAPI Web 服务 - 替代 Gradio UI，无锁并发体验。

特点：
- 流式聊天用 SSE（Server-Sent Events）；前端用 EventSource 接收。
- 长任务（如对话）放在线程池里跑，主事件循环始终空闲，
  思考期间点任意按钮/上传/切用户都能立即响应。
- 上传支持 multipart 多文件 + 文件夹（前端用 webkitdirectory）。
- 单页静态 UI（static/index.html）+ 异步 fetch API。
"""
from __future__ import annotations

import asyncio
import copy
import hashlib
import json
import logging
import math
import os
import re
import shutil
import tempfile
import threading
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, AsyncGenerator

import pandas as pd
from fastapi import FastAPI, File, Form, Header, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles

from agent_core import StorageAgent
from capacity_gpt55 import (
    DEFAULT_NA_CELL_COST_YUAN_PER_WH,
    DEFAULT_STORAGE_COST_BASIS,
    DEFAULT_STORAGE_SYSTEM_COST_YUAN_PER_KWH,
    analyze_capacity_with_bill_method,
    build_optimal_config_from_record,
)
from config import AgentConfig
from hier_memory import HierarchicalMemory, safe_user_id
from storage_document_analysis import analyze_storage_documents

logger = logging.getLogger(__name__)

DATA_ROOT = Path(os.environ.get("CHUNENG_DATA_ROOT", "/var/lib/chuneng-agent")).resolve()
INPUT_DIR = (DATA_ROOT / "input").resolve()
INPUT_DIR.mkdir(parents=True, exist_ok=True)
STATIC_DIR = Path(__file__).parent / "static"

SUPPORTED_UPLOAD_EXTS = {".pdf", ".docx", ".doc", ".xlsx", ".xls",
                         ".png", ".jpg", ".jpeg", ".csv", ".txt"}
SUPPORTED_KB_EXTS = {".txt", ".md", ".markdown", ".pdf", ".docx", ".xlsx", ".csv"}
SUPPORTED_BILL_EXTS = {".pdf", ".docx", ".doc", ".xlsx", ".xls",
                       ".png", ".jpg", ".jpeg", ".csv", ".txt"}
REPORT_FILE_EXTS = {".zip", ".docx", ".md", ".png"}
BILL_COLUMNS = [
    "月份", "总电量(kWh)", "尖峰电量(kWh)", "高峰电量(kWh)", "平段电量(kWh)",
    "谷段电量(kWh)", "最大需量(kW)", "合同容量(kVA)", "总电费(元)",
    "电量电费(元)", "需量电费(元)", "容量电费(元)", "功率因数",
    "尖峰电价(元/kWh)", "高峰电价(元/kWh)", "平段电价(元/kWh)",
    "谷段电价(元/kWh)", "需量电价(元/kW·月)",
]
REVENUE_PARAM_HISTORY_LIMIT = 40
DEFAULT_STORAGE_SYSTEM_COST_YUAN_PER_WH = DEFAULT_STORAGE_SYSTEM_COST_YUAN_PER_KWH / 1000
DEFAULT_PCS_EMS_BMS_COST_YUAN_PER_WH = 0.025
DEFAULT_GRID_CONNECTION_COST_PER_KW = 10
DEFAULT_CIVIL_FIRE_COST_PER_KW = 0
DEFAULT_INCLUDED_COST_RATE = 0.0


# ======================================================================
# AgentManager（多用户隔离）
# ======================================================================
class AgentManager:
    def __init__(self, config: AgentConfig, **kwargs):
        self.config = config
        self.kwargs = kwargs
        self.agents: dict[str, StorageAgent] = {}
        self._lock = threading.Lock()

    def get(self, user_id: str) -> StorageAgent:
        uid = safe_user_id(user_id)
        with self._lock:
            if uid not in self.agents:
                logger.info("创建新 Agent: user_id=%s", uid)
                self.agents[uid] = StorageAgent(
                    config=self.config, user_id=uid, **self.kwargs
                )
                try:
                    _load_persisted_state(self.agents[uid])
                except Exception:
                    logger.exception("加载本地历史状态失败: user_id=%s", uid)
            return self.agents[uid]

    def list_users(self) -> list[str]:
        try:
            users = HierarchicalMemory.list_users()
        except Exception:
            users = []
        try:
            state_root = Path(self.config.output_dir) / "web_state"
            if state_root.exists():
                for path in state_root.iterdir():
                    if path.is_dir() and path.name not in users:
                        users.append(path.name)
        except Exception:
            pass
        # 合并已加载的 user_id
        with self._lock:
            for uid in self.agents.keys():
                if uid not in users:
                    users.append(uid)
        if not users:
            users = ["main"]
        return sorted(set(users))


# ======================================================================
# 工具函数
# ======================================================================
def _state_summary(agent: StorageAgent) -> dict:
    s = agent.state
    has_data = s.electricity_df is not None and not s.electricity_df.empty
    persisted_capacity = _read_json(_state_path(agent, "capacity_analysis.json")) or {}
    persisted_revenue = _read_json(_state_path(agent, "revenue_analysis.json")) or {}
    capacity_rows = persisted_capacity.get("results") or getattr(s, "capacity_analysis", None) or []
    return {
        "user_id": agent.user_id,
        "has_data": has_data,
        "rows": int(len(s.electricity_df)) if has_data else 0,
        "has_optimization": s.optimal_config is not None or bool(capacity_rows),
        "has_revenue": s.revenue_report is not None or getattr(s, "revenue_model", None) is not None or bool(persisted_revenue),
        "has_investor": s.investor_report is not None,
        "has_md_report": s.md_report is not None,
        "input_files": sorted([p.name for p in INPUT_DIR.glob("*") if p.is_file()]),
        "tools_count": len(agent.registry.all()),
        "react": agent.enable_react,
    }


def _memory_summary(agent: StorageAgent) -> dict:
    if agent.state.memory is None:
        return {"enabled": False}
    try:
        return {
            "enabled": True,
            "stats": agent.state.memory.stats(),
            "facts": agent.state.memory.list_facts(),
        }
    except Exception as e:
        return {"enabled": True, "error": str(e)}


def _kb_summary(agent: StorageAgent) -> dict:
    kb = agent.state.kb
    if kb is None or not getattr(kb, "is_ready", False):
        return {"enabled": False}
    try:
        stats = kb.stats() if hasattr(kb, "stats") else {}
        docs = kb.list_documents() if hasattr(kb, "list_documents") else []
        return {"enabled": True, "stats": stats, "documents": docs}
    except Exception as e:
        return {"enabled": True, "error": str(e)}


def _save_upload(file: UploadFile, allowed_exts: set[str] | None) -> tuple[Path | None, str]:
    """保存上传文件到 INPUT_DIR；同名自动加序号。返回 (Path, msg)。"""
    raw_name = Path(file.filename or "").name
    if not raw_name:
        return None, "文件名为空"
    suffix = Path(raw_name).suffix.lower()
    if allowed_exts is not None and suffix and suffix not in allowed_exts:
        return None, f"不支持的扩展名: {suffix}"
    dst = INPUT_DIR / raw_name
    if dst.exists():
        stem, suf = dst.stem, dst.suffix
        i = 1
        while dst.exists():
            dst = INPUT_DIR / f"{stem}_{i}{suf}"
            i += 1
    try:
        with open(dst, "wb") as fp:
            shutil.copyfileobj(file.file, fp)
    except Exception as e:
        return None, f"写入失败: {e}"
    return dst, "ok"


def _list_bill_files() -> list[str]:
    return sorted(
        p.name for p in INPUT_DIR.glob("*")
        if p.is_file() and p.suffix.lower() in SUPPORTED_BILL_EXTS
    )


def _resolve_input_file(name: str) -> Path:
    raw_name = Path(name or "").name
    if not raw_name:
        raise HTTPException(400, "文件名为空")
    target = (INPUT_DIR / raw_name).resolve()
    if target.parent != INPUT_DIR.resolve():
        raise HTTPException(400, "非法路径")
    if not target.exists() or not target.is_file():
        raise HTTPException(404, f"文件不存在: {raw_name}")
    if target.suffix.lower() not in SUPPORTED_BILL_EXTS:
        raise HTTPException(400, f"不支持的账单文件类型: {target.suffix}")
    return target


def _normalize_month(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    m = re.search(r"(20\d{2})\D{0,3}([01]?\d)", text)
    if m:
        month = max(1, min(12, int(m.group(2))))
        return f"{m.group(1)}-{month:02d}"
    m = re.search(r"(20\d{2})([01]\d)", text)
    if m:
        month = max(1, min(12, int(m.group(2))))
        return f"{m.group(1)}-{month:02d}"
    return text


def _bills_to_dataframe(bills: list[Any]) -> pd.DataFrame:
    records = []
    for b in bills:
        records.append({
            "月份": _normalize_month(getattr(b, "month", "")),
            "总电量(kWh)": getattr(b, "total_kwh", 0),
            "尖峰电量(kWh)": getattr(b, "peak_kwh", 0),
            "高峰电量(kWh)": getattr(b, "high_kwh", 0),
            "平段电量(kWh)": getattr(b, "flat_kwh", 0),
            "谷段电量(kWh)": getattr(b, "valley_kwh", 0),
            "最大需量(kW)": getattr(b, "max_demand_kw", 0),
            "合同容量(kVA)": getattr(b, "contract_capacity_kva", 0),
            "总电费(元)": getattr(b, "total_amount", 0),
            "电量电费(元)": getattr(b, "energy_charge", 0),
            "需量电费(元)": getattr(b, "demand_charge", 0),
            "容量电费(元)": getattr(b, "capacity_charge", 0),
            "功率因数": getattr(b, "power_factor", 0),
            "尖峰电价(元/kWh)": getattr(b, "peak_price", 0),
            "高峰电价(元/kWh)": getattr(b, "high_price", 0),
            "平段电价(元/kWh)": getattr(b, "flat_price", 0),
            "谷段电价(元/kWh)": getattr(b, "valley_price", 0),
            "需量电价(元/kW·月)": getattr(b, "demand_price", 0),
        })
    return pd.DataFrame(records)


def _normalize_bill_df(df: pd.DataFrame | None) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame(columns=[*BILL_COLUMNS, "平均电价(元/kWh)"])

    out = df.copy()
    for col in BILL_COLUMNS:
        if col not in out.columns:
            out[col] = "" if col == "月份" else 0

    out["月份"] = out["月份"].map(_normalize_month)
    for col in BILL_COLUMNS:
        if col == "月份":
            continue
        out[col] = pd.to_numeric(out[col], errors="coerce").fillna(0)

    out = out[(out["总电量(kWh)"] > 0) | (out["总电费(元)"] > 0)].copy()
    if out.empty:
        return pd.DataFrame(columns=[*BILL_COLUMNS, "平均电价(元/kWh)"])

    out["平均电价(元/kWh)"] = 0.0
    mask = out["总电量(kWh)"] > 0
    out.loc[mask, "平均电价(元/kWh)"] = out.loc[mask, "总电费(元)"] / out.loc[mask, "总电量(kWh)"]
    out = out[[*BILL_COLUMNS, "平均电价(元/kWh)"]]
    return out.sort_values("月份", kind="stable").reset_index(drop=True)


def _merge_bill_df(existing: pd.DataFrame | None, incoming: pd.DataFrame | None) -> pd.DataFrame:
    old = _normalize_bill_df(existing)
    new = _normalize_bill_df(incoming)
    if old.empty:
        return new
    if new.empty:
        return old
    combined = pd.concat([old, new], ignore_index=True)
    combined["_completeness"] = combined.apply(
        lambda row: sum(
            1 for col in BILL_COLUMNS
            if col != "月份" and _safe_float(row.get(col), 6, 0) not in (None, 0)
        ),
        axis=1,
    )
    combined["_row_order"] = range(len(combined))
    combined = combined.sort_values(["月份", "_completeness", "_row_order"], kind="stable")
    combined = combined.drop_duplicates(subset=["月份"], keep="last")
    combined = combined.drop(columns=["_completeness", "_row_order"], errors="ignore")
    return _normalize_bill_df(combined)


def _round_float(value: Any, digits: int = 2) -> float:
    try:
        v = float(value)
    except Exception:
        return 0.0
    if not math.isfinite(v):
        return 0.0
    return round(v, digits)


def _safe_float(value: Any, digits: int = 2, default: float | None = 0.0) -> float | None:
    try:
        v = float(value)
    except Exception:
        return default
    if not math.isfinite(v):
        return default
    return round(v, digits)


def _effective_user_id(user_id: str | None = None, auth_user: str | None = None) -> str:
    explicit = str(user_id or "").strip()
    if explicit:
        return safe_user_id(explicit)
    return safe_user_id(str(auth_user or "").strip() or "main")


def _user_state_dir(agent: StorageAgent | str) -> Path:
    user_id = agent.user_id if hasattr(agent, "user_id") else str(agent)
    base = Path(getattr(agent.state.config, "output_dir", "output")) if hasattr(agent, "state") else Path("output")
    path = base / "web_state" / safe_user_id(user_id)
    path.mkdir(parents=True, exist_ok=True)
    return path


def _json_default(value: Any):
    if isinstance(value, pd.DataFrame):
        return value.where(pd.notna(value), None).to_dict(orient="records")
    if hasattr(value, "__dict__"):
        return {
            key: val for key, val in vars(value).items()
            if not key.startswith("_")
        }
    return str(value)


def _state_path(agent: StorageAgent, name: str) -> Path:
    return _user_state_dir(agent) / name


def _write_json(path: Path, data: dict):
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(json.dumps(data, ensure_ascii=False, indent=2, default=_json_default), encoding="utf-8")
    tmp.replace(path)


def _read_json(path: Path) -> dict | None:
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        logger.exception("读取本地状态失败: %s", path)
        return None


def _persist_bills(agent: StorageAgent, df: pd.DataFrame, *, files: list[str] | None = None,
                   parser: str = "", msg: str = "") -> dict:
    existing = _read_json(_state_path(agent, "bills.json")) or {}
    if files is None:
        files = existing.get("files") or []
    if not parser:
        parser = existing.get("parser") or ""
    if not msg:
        msg = existing.get("msg") or ""
    payload = _bill_payload(df, files=files, parser=parser, msg=msg)
    payload["saved_at"] = datetime.now().isoformat(timespec="seconds")
    _write_json(_state_path(agent, "bills.json"), payload)
    return payload


def _load_bills(agent: StorageAgent) -> dict | None:
    data = _read_json(_state_path(agent, "bills.json"))
    if not data:
        return None
    records = data.get("records") or []
    if records:
        df = _normalize_bill_df(pd.DataFrame(records))
        if not df.empty:
            agent.state.electricity_df = df
            payload = _bill_payload(
                df,
                files=data.get("files") or [],
                parser=data.get("parser") or "",
                msg=data.get("msg") or "已加载本地历史账单数据",
            )
            payload["saved_at"] = data.get("saved_at")
            return payload
    return None


def _persist_capacity(agent: StorageAgent, data: dict) -> dict:
    out = dict(data)
    out["saved_at"] = datetime.now().isoformat(timespec="seconds")
    _write_json(_state_path(agent, "capacity_analysis.json"), out)
    return out


def _load_capacity(agent: StorageAgent) -> dict | None:
    return _read_json(_state_path(agent, "capacity_analysis.json"))


def _persist_revenue(agent: StorageAgent, data: dict) -> dict:
    out = dict(data)
    out["saved_at"] = datetime.now().isoformat(timespec="seconds")
    _write_json(_state_path(agent, "revenue_analysis.json"), out)
    _persist_revenue_param_history(agent, out)
    return out


def _load_revenue(agent: StorageAgent) -> dict | None:
    return _read_json(_state_path(agent, "revenue_analysis.json"))


def _load_revenue_param_history(agent: StorageAgent) -> dict:
    data = _read_json(_state_path(agent, "revenue_params_history.json")) or {}
    items = data.get("items") if isinstance(data, dict) else []
    return {
        "ok": True,
        "version": 1,
        "items": items if isinstance(items, list) else [],
    }


def _revenue_param_signature(params: dict) -> str:
    ignored = {"user_id", "model_source", "model_version", "saved_at"}
    stable = {k: v for k, v in params.items() if k not in ignored}
    raw = json.dumps(stable, ensure_ascii=False, sort_keys=True, default=_json_default)
    return hashlib.sha1(raw.encode("utf-8")).hexdigest()[:16]


def _revenue_param_label(params: dict, summary: dict | None = None) -> str:
    power = _safe_float(params.get("power_kw"), 2) or 0
    capacity = _safe_float(params.get("capacity_kwh"), 2) or 0
    duration = _safe_float(params.get("duration_hours"), 2) or (capacity / power if power else 0)
    demand_unit = _safe_float(params.get("demand_revenue_per_kw_year"), 2) or 0
    demand = _safe_float(params.get("demand_revenue"), 2) or _safe_float(demand_unit * power, 2) or 0
    spread = _safe_float((summary or {}).get("spread_yuan_per_kwh"), 4, None)
    label = f"{power:g}kW / {capacity:g}kWh / {duration:g}h"
    if demand > 0:
        label += f" · 需量{demand / 10000:g}万元/年"
    if spread is not None:
        label += f" · 价差{spread:g}元/kWh"
    return label


def _persist_revenue_param_history(agent: StorageAgent, revenue: dict) -> dict:
    params = revenue.get("params") if isinstance(revenue, dict) else None
    if not isinstance(params, dict) or not params:
        return _load_revenue_param_history(agent)

    saved_at = revenue.get("saved_at") or datetime.now().isoformat(timespec="seconds")
    summary = revenue.get("summary") if isinstance(revenue.get("summary"), dict) else {}
    signature = _revenue_param_signature(params)
    current = _load_revenue_param_history(agent)
    items = [item for item in current.get("items", []) if item.get("signature") != signature]
    item = {
        "id": f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{signature}",
        "signature": signature,
        "saved_at": saved_at,
        "source": revenue.get("params_source") or "manual",
        "label": _revenue_param_label(params, summary),
        "params": dict(params),
        "summary": {
            "power_kw": summary.get("power_kw"),
            "capacity_kwh": summary.get("capacity_kwh"),
            "duration_hours": summary.get("duration_hours"),
            "spread_yuan_per_kwh": summary.get("spread_yuan_per_kwh"),
            "first_year_gross_revenue_yuan": summary.get("first_year_gross_revenue_yuan"),
            "static_payback_years": summary.get("static_payback_years"),
            "project_irr_percent": summary.get("project_irr_percent"),
        },
    }
    if item["source"] == "capacity_sync":
        item["label"] = f"推荐配置 · {item['label']}"
    items.insert(0, item)
    payload = {"ok": True, "version": 1, "items": items[:REVENUE_PARAM_HISTORY_LIMIT]}
    _write_json(_state_path(agent, "revenue_params_history.json"), payload)
    return payload


def _best_capacity_record(agent: StorageAgent) -> dict:
    capacity = _load_capacity(agent) or {}
    best = capacity.get("best") or {}
    if best:
        return best
    rows = capacity.get("results") or getattr(agent.state, "capacity_analysis", None) or []
    if rows:
        return next((row for row in rows if row.get("is_best")), rows[0])
    return {}


def _revenue_matches_capacity(revenue: dict | None, capacity: dict | None) -> bool:
    if not revenue or not capacity:
        return False
    params = revenue.get("params") or {}
    summary = revenue.get("summary") or {}
    best = capacity.get("best") or {}
    if not best:
        rows = capacity.get("results") or []
        best = next((row for row in rows if row.get("is_best")), rows[0]) if rows else {}
    if not best:
        return True
    expected_capacity = _safe_float(best.get("battery_capacity_kwh"), 2, None)
    expected_power = _safe_float(best.get("inverter_power_kw"), 2, None)
    expected_duration = _safe_float(best.get("duration_hours"), 2, None)
    actual_capacity = _safe_float(params.get("capacity_kwh") or summary.get("capacity_kwh"), 2, None)
    actual_power = _safe_float(params.get("power_kw") or summary.get("power_kw"), 2, None)
    actual_duration = _safe_float(params.get("duration_hours") or summary.get("duration_hours"), 2, None)
    return (
        expected_capacity == actual_capacity
        and expected_power == actual_power
        and (expected_duration is None or actual_duration == expected_duration)
    )


def _revenue_params_for_current_capacity(agent: StorageAgent, user_id: str = "main") -> dict:
    saved = _load_revenue(agent) or {}
    params = dict(saved.get("params") or {}) if isinstance(saved.get("params"), dict) else {}
    for key in ("power_kw", "duration_hours", "capacity_kwh"):
        params.pop(key, None)
    params["user_id"] = user_id
    return params


def _refresh_revenue_for_current_capacity(agent: StorageAgent, user_id: str = "main") -> dict:
    params = _revenue_params_for_current_capacity(agent, user_id)
    params["_sync_to_capacity"] = True
    params["battery_unit_cost"] = DEFAULT_NA_CELL_COST_YUAN_PER_WH
    params["pcs_ems_bms_unit_cost"] = DEFAULT_PCS_EMS_BMS_COST_YUAN_PER_WH
    params["grid_connection_cost_per_kw"] = DEFAULT_GRID_CONNECTION_COST_PER_KW
    params["civil_fire_cost_per_kw"] = DEFAULT_CIVIL_FIRE_COST_PER_KW
    params["cell_cost_yuan_per_wh"] = DEFAULT_NA_CELL_COST_YUAN_PER_WH
    params["system_unit_cost_yuan_per_wh"] = DEFAULT_STORAGE_SYSTEM_COST_YUAN_PER_WH
    params["cost_basis"] = DEFAULT_STORAGE_COST_BASIS
    params["design_supervision_rate"] = DEFAULT_INCLUDED_COST_RATE
    params["contingency_rate"] = DEFAULT_INCLUDED_COST_RATE
    params["construction_interest_rate"] = DEFAULT_INCLUDED_COST_RATE
    return _compute_revenue_model(agent, params)


def _load_persisted_state(agent: StorageAgent):
    _load_bills(agent)
    capacity = _load_capacity(agent)
    if capacity and capacity.get("results"):
        agent.state.capacity_analysis = capacity.get("results")
        try:
            best = capacity.get("best") or next((row for row in capacity.get("results") or [] if row.get("is_best")), None)
            if best:
                agent.state.optimal_config = build_optimal_config_from_record(best, agent.state.config.storage_config)
        except Exception:
            logger.exception("加载本地容量推荐配置失败: user_id=%s", agent.user_id)
    revenue = _load_revenue(agent)
    if revenue:
        agent.state.revenue_model = revenue


def _persist_runtime_state(agent: StorageAgent):
    df = _normalize_bill_df(agent.state.electricity_df)
    if not df.empty:
        _persist_bills(agent, df, msg="已保存 Agent 当前账单数据")

    capacity = getattr(agent.state, "capacity_analysis", None)
    if capacity:
        saved = _load_capacity(agent) or {}
        if not saved.get("results"):
            _persist_capacity(agent, {
                "ok": True,
                "msg": "已保存 Agent 当前容量分析结果",
                "candidate_count": len(capacity),
                "positive_count": sum(1 for row in capacity if (row.get("annual_revenue_yuan") or 0) > 0),
                "best": next((row for row in capacity if row.get("is_best")), capacity[0]),
                "results": capacity,
            })


def _clear_persisted_state(agent: StorageAgent):
    path = _user_state_dir(agent)
    if path.exists():
        shutil.rmtree(path)


def _bill_payload(df: pd.DataFrame | None, *, files: list[str] | None = None,
                  parser: str = "", msg: str = "") -> dict:
    norm = _normalize_bill_df(df)
    if norm.empty:
        return {
            "ok": True,
            "msg": msg,
            "parser": parser,
            "files": files or [],
            "available_files": _list_bill_files(),
            "records": [],
            "columns": [*BILL_COLUMNS, "平均电价(元/kWh)"],
            "summary": {
                "row_count": 0,
                "month_count": 0,
                "total_kwh": 0,
                "total_amount": 0,
                "avg_unit_price": 0,
                "max_demand_kw": 0,
                "tou": {"peak": 0, "high": 0, "flat": 0, "valley": 0},
            },
        }

    summary = {
        "row_count": int(len(norm)),
        "month_count": int(norm["月份"].nunique()),
        "start_month": str(norm["月份"].iloc[0]),
        "end_month": str(norm["月份"].iloc[-1]),
        "total_kwh": _round_float(norm["总电量(kWh)"].sum(), 2),
        "total_amount": _round_float(norm["总电费(元)"].sum(), 2),
        "avg_unit_price": _round_float(
            norm["总电费(元)"].sum() / norm["总电量(kWh)"].sum()
            if norm["总电量(kWh)"].sum() > 0 else 0,
            4,
        ),
        "max_demand_kw": _round_float(norm["最大需量(kW)"].max(), 2),
        "tou": {
            "peak": _round_float(norm["尖峰电量(kWh)"].sum(), 2),
            "high": _round_float(norm["高峰电量(kWh)"].sum(), 2),
            "flat": _round_float(norm["平段电量(kWh)"].sum(), 2),
            "valley": _round_float(norm["谷段电量(kWh)"].sum(), 2),
        },
        "prices": {
            "peak": _round_float(norm.loc[norm["尖峰电价(元/kWh)"] > 0, "尖峰电价(元/kWh)"].mean(), 4),
            "high": _round_float(norm.loc[norm["高峰电价(元/kWh)"] > 0, "高峰电价(元/kWh)"].mean(), 4),
            "flat": _round_float(norm.loc[norm["平段电价(元/kWh)"] > 0, "平段电价(元/kWh)"].mean(), 4),
            "valley": _round_float(norm.loc[norm["谷段电价(元/kWh)"] > 0, "谷段电价(元/kWh)"].mean(), 4),
            "demand": _round_float(norm.loc[norm["需量电价(元/kW·月)"] > 0, "需量电价(元/kW·月)"].mean(), 4),
        },
    }
    records = norm.where(pd.notna(norm), None).to_dict(orient="records")
    return {
        "ok": True,
        "msg": msg,
        "parser": parser,
        "files": files or [],
        "available_files": _list_bill_files(),
        "records": records,
        "columns": [*BILL_COLUMNS, "平均电价(元/kWh)"],
        "summary": summary,
    }


def _parse_bill_paths_sync(agent: StorageAgent, paths: list[Path], mode: str) -> tuple[pd.DataFrame, str]:
    parser_used = "rules"
    df = pd.DataFrame()

    if mode != "rules" and agent.state.llm_parser is not None:
        bills = agent.state.llm_parser.parse_batch(paths)
        if bills:
            df = _bills_to_dataframe(bills)
            parser_used = "llm"

    if df.empty:
        parsed = agent.state.parser.parse_batch(paths)
        df = agent.state.extractor.extract_from_parsed(parsed)
        parser_used = "rules"

    df = _normalize_bill_df(df)
    if df.empty:
        return df, parser_used

    existing = _normalize_bill_df(agent.state.electricity_df)
    combined = _merge_bill_df(existing, df)
    saved = _read_json(_state_path(agent, "bills.json")) or {}
    merged_files = sorted(set((saved.get("files") or []) + [p.name for p in paths]))
    agent.state.electricity_df = combined
    agent.state.optimal_config = None
    agent.state.revenue_report = None
    agent.state.investor_report = None
    agent.state.md_report = None
    _persist_bills(agent, combined, files=merged_files, parser=parser_used,
                   msg=f"已解析 {len(paths)} 个文件，新增/更新 {len(df)} 条账单记录；当前共 {len(combined)} 条")
    return df, parser_used


def _parse_single_bill_path_sync(agent: StorageAgent, path: Path, mode: str) -> tuple[pd.DataFrame, str]:
    return _parse_bill_paths_sync(agent, [path], mode)


def _analyze_storage_docs_sync(agent: StorageAgent, paths: list[Path], topic: str,
                               index_to_kb: bool = True, use_llm: bool = False) -> dict:
    return analyze_storage_documents(
        paths,
        agent.state.parser,
        llm_client=getattr(agent.state, "llm_client", None),
        kb=getattr(agent.state, "kb", None),
        topic=topic,
        index_to_kb=index_to_kb,
        use_llm=use_llm,
    )


def _set_bill_state(agent: StorageAgent, df: pd.DataFrame, *, files: list[str] | None = None,
                    parser: str = "", msg: str = "", merge: bool = False) -> pd.DataFrame:
    df = _merge_bill_df(agent.state.electricity_df, df) if merge else _normalize_bill_df(df)
    if merge:
        saved = _read_json(_state_path(agent, "bills.json")) or {}
        files = sorted(set((saved.get("files") or []) + (files or [])))
    agent.state.electricity_df = df
    agent.state.optimal_config = None
    agent.state.revenue_report = None
    agent.state.investor_report = None
    agent.state.md_report = None
    _persist_bills(agent, df, files=files, parser=parser, msg=msg or f"已保存 {len(df)} 条账单记录")
    return df


def _nice_capacity(value: float) -> float:
    if value >= 5000:
        step = 500
    elif value >= 1000:
        step = 100
    elif value >= 200:
        step = 50
    else:
        step = 10
    return float(max(step, round(value / step) * step))


def _capacity_score(cfg: Any) -> float:
    investment = float(getattr(cfg, "total_investment", 0) or 0)
    annual_revenue = float(getattr(cfg, "annual_revenue", 0) or 0)
    payback = float(getattr(cfg, "simple_payback_years", math.inf) or math.inf)
    irr = float(getattr(cfg, "irr", 0) or 0)
    npv = float(getattr(cfg, "npv", 0) or 0)
    if annual_revenue <= 0 or not math.isfinite(payback):
        return -1_000_000 + annual_revenue
    npv_ratio = npv / investment if investment > 0 else 0
    return irr * 100 + npv_ratio * 30 - payback * 2


def _capacity_record(name: str, cfg: Any, score: float) -> dict:
    payback = _safe_float(getattr(cfg, "simple_payback_years", None), 2, None)
    return {
        "name": name,
        "battery_capacity_kwh": _safe_float(getattr(cfg, "battery_capacity_kwh", 0), 2),
        "inverter_power_kw": _safe_float(getattr(cfg, "inverter_power_kw", 0), 2),
        "duration_hours": _safe_float(getattr(cfg, "duration_hours", 0), 2),
        "daily_charge_kwh": _safe_float(getattr(cfg, "daily_charge_kwh", 0), 2),
        "daily_discharge_kwh": _safe_float(getattr(cfg, "daily_discharge_kwh", 0), 2),
        "total_investment_yuan": _safe_float(getattr(cfg, "total_investment", 0), 2),
        "annual_savings_yuan": _safe_float(getattr(cfg, "annual_savings", 0), 2),
        "annual_revenue_yuan": _safe_float(getattr(cfg, "annual_revenue", 0), 2),
        "payback_years": payback,
        "npv_yuan": _safe_float(getattr(cfg, "npv", 0), 2),
        "irr": _safe_float(getattr(cfg, "irr", 0), 4),
        "irr_percent": _safe_float(float(getattr(cfg, "irr", 0) or 0) * 100, 2),
        "lcoe_yuan_per_kwh": _safe_float(getattr(cfg, "lcoe", 0), 4),
        "charge_window": f"{getattr(cfg, 'charge_start_hour', 0)}:00-{getattr(cfg, 'charge_end_hour', 0)}:00",
        "discharge_window": f"{getattr(cfg, 'discharge_start_hour', 0)}:00-{getattr(cfg, 'discharge_end_hour', 0)}:00",
        "score": _safe_float(score, 4),
    }


def _capacity_analysis_sync(agent: StorageAgent, req: dict | None = None) -> dict:
    req = req or {}
    df = _normalize_bill_df(agent.state.electricity_df)
    if df.empty:
        _load_bills(agent)
        df = _normalize_bill_df(agent.state.electricity_df)
    if df.empty:
        raise HTTPException(400, "请先解析账单数据，再进行储能容量分析")

    analyzer = agent.state.analyzer
    storage = agent.state.config.storage_config
    method = str(req.get("method") or "gpt55").strip().lower()
    if method not in {"legacy", "classic", "old"}:
        try:
            result = analyze_capacity_with_bill_method(
                df,
                agent.state.config.rate_config,
                storage,
                req,
            )
        except ValueError as e:
            raise HTTPException(422, str(e))

        best_record = result.get("best") or {}
        best_cfg = build_optimal_config_from_record(best_record, storage)
        agent.state.optimal_config = best_cfg
        agent.state.revenue_report = analyzer.analyze(best_cfg, df)
        agent.state.investor_report = None
        agent.state.md_report = None
        agent.state.capacity_analysis = result.get("results") or []
        persisted = _persist_capacity(agent, result)
        try:
            persisted["revenue_model"] = _refresh_revenue_for_current_capacity(agent, req.get("user_id", agent.user_id))
        except Exception as e:
            logger.exception("容量分析后同步收益测算失败: %s", e)
            persisted["revenue_sync_error"] = str(e)
        return persisted

    optimizer = agent.state.optimizer
    load_profile = optimizer._analyze_load_profile(df)
    base_cfg = optimizer.optimize(df)

    raw_capacities = req.get("capacities_kwh") or req.get("capacities") or []
    capacities: list[float] = []
    if isinstance(raw_capacities, (int, float, str)):
        raw_capacities = [raw_capacities]
    for value in raw_capacities:
        try:
            v = float(value)
        except Exception:
            continue
        if v > 0:
            capacities.append(_nice_capacity(v))

    if not capacities:
        base = max(float(base_cfg.battery_capacity_kwh or 0), float(storage.min_capacity_kwh or 0))
        multipliers = [0.5, 0.75, 1.0, 1.25, 1.5, 1.75, 2.0]
        capacities = [_nice_capacity(base * m) for m in multipliers]
        capacities.append(_nice_capacity(load_profile.get("daily_valley", 0) * 0.8))
        capacities.append(_nice_capacity((load_profile.get("daily_peak", 0) + load_profile.get("daily_high", 0)) * 0.6))

    capacities = sorted({
        max(float(storage.min_capacity_kwh), min(float(storage.max_capacity_kwh), c))
        for c in capacities if c > 0
    })

    raw_durations = req.get("durations_hours") or req.get("durations") or [2, 3, 4]
    if isinstance(raw_durations, (int, float, str)):
        raw_durations = [raw_durations]
    durations: list[float] = []
    for value in raw_durations:
        try:
            v = float(value)
        except Exception:
            continue
        if v > 0:
            durations.append(v)
    durations = sorted(set(durations or [2, 3, 4]))

    peak_limit = float(load_profile.get("peak_load_kw") or 0) * 0.5
    max_power = min(float(storage.max_power_kw), peak_limit) if peak_limit > 0 else float(storage.max_power_kw)
    min_power = float(storage.min_power_kw)

    candidates: list[tuple[int, str, Any, float]] = []
    seen: set[tuple[float, float]] = set()

    def add_candidate(name: str, capacity_kwh: float, power_kw: float):
        if capacity_kwh <= 0 or power_kw <= 0:
            return
        capacity_kwh = max(float(storage.min_capacity_kwh), min(float(storage.max_capacity_kwh), capacity_kwh))
        power_kw = max(min_power, min(max_power, power_kw))
        key = (round(capacity_kwh, 2), round(power_kw, 2))
        if key in seen:
            return
        seen.add(key)
        strategy = optimizer._optimize_charge_strategy(load_profile, capacity_kwh, power_kw)
        cfg = optimizer._calculate_economics(capacity_kwh, power_kw, strategy, df, load_profile)
        if name != "优化器推荐":
            name = f"{cfg.battery_capacity_kwh:g}kWh / {cfg.duration_hours:g}h"
        candidates.append((len(candidates), name, cfg, _capacity_score(cfg)))

    add_candidate("优化器推荐", float(base_cfg.battery_capacity_kwh), float(base_cfg.inverter_power_kw))
    for capacity in capacities:
        for duration in durations:
            add_candidate(f"{capacity:g}kWh / {duration:g}h", capacity, capacity / duration)

    if not candidates:
        raise HTTPException(422, "未能生成有效储能容量组合")

    records = [_capacity_record(name, cfg, score) for _, name, cfg, score in candidates]
    records.sort(
        key=lambda r: (
            r["annual_revenue_yuan"] > 0,
            r["score"] if r["score"] is not None else -1_000_000,
            r["npv_yuan"] if r["npv_yuan"] is not None else -1_000_000,
        ),
        reverse=True,
    )
    for idx, rec in enumerate(records, start=1):
        rec["rank"] = idx
        rec["is_best"] = idx == 1

    best_record = records[0]
    best_idx = next(
        idx for idx, name, cfg, score in candidates
        if name == best_record["name"]
        and _safe_float(getattr(cfg, "battery_capacity_kwh", 0), 2) == best_record["battery_capacity_kwh"]
        and _safe_float(getattr(cfg, "inverter_power_kw", 0), 2) == best_record["inverter_power_kw"]
    )
    best_cfg = candidates[best_idx][2]
    agent.state.optimal_config = best_cfg
    agent.state.revenue_report = analyzer.analyze(best_cfg, df)
    agent.state.investor_report = None
    agent.state.md_report = None
    agent.state.capacity_analysis = records

    positive_count = sum(1 for r in records if (r.get("annual_revenue_yuan") or 0) > 0)
    result = {
        "ok": True,
        "msg": f"已分析 {len(records)} 个储能容量组合，推荐 {best_record['battery_capacity_kwh']:g}kWh / {best_record['inverter_power_kw']:g}kW",
        "candidate_count": len(records),
        "positive_count": positive_count,
        "scoring_basis": "综合评分 = IRR*100 + NPV/投资额*30 - 回收期*2；仅净收益为正的组合优先。",
        "load_profile": {
            "daily_kwh": _safe_float(load_profile.get("daily_kwh"), 2),
            "daily_peak_high_kwh": _safe_float(
                load_profile.get("daily_peak", 0) + load_profile.get("daily_high", 0), 2
            ),
            "daily_valley_kwh": _safe_float(load_profile.get("daily_valley"), 2),
            "max_demand_kw": _safe_float(load_profile.get("max_demand_kw"), 2),
        },
        "best": best_record,
        "results": records,
    }
    persisted = _persist_capacity(agent, result)
    try:
        persisted["revenue_model"] = _refresh_revenue_for_current_capacity(agent, req.get("user_id", agent.user_id))
    except Exception as e:
        logger.exception("容量分析后同步收益测算失败: %s", e)
        persisted["revenue_sync_error"] = str(e)
    return persisted


REVENUE_MODEL_VERSION = "excel-11-20260509-v2"
REVENUE_TEMPLATE_PATH = "/opt/download-hub/storage/11_20260509160117.xlsx"


def _default_revenue_params(agent: StorageAgent) -> dict:
    storage = agent.state.config.storage_config
    cfg = agent.state.optimal_config
    if cfg is None:
        best = _best_capacity_record(agent)
        if best:
            capacity = float(best.get("battery_capacity_kwh") or getattr(storage, "min_capacity_kwh", 1000) or 1000)
            power = float(best.get("inverter_power_kw") or max(float(getattr(storage, "min_power_kw", 500) or 500), capacity / 2))
            duration = float(best.get("duration_hours") or (capacity / power if power else 2))
        else:
            power = 1250.0
            duration = 4.0
            capacity = 4480.0
    else:
        capacity = float(getattr(cfg, "battery_capacity_kwh", 0) or 0)
        power = float(getattr(cfg, "inverter_power_kw", 0) or 0)
        duration = float(getattr(cfg, "duration_hours", 0) or 0)
    return {
        "model_version": REVENUE_MODEL_VERSION,
        "model_source": REVENUE_TEMPLATE_PATH,
        "project_name": "钠离子储能项目测算",
        "cost_basis": DEFAULT_STORAGE_COST_BASIS,
        "cell_cost_yuan_per_wh": DEFAULT_NA_CELL_COST_YUAN_PER_WH,
        "system_unit_cost_yuan_per_wh": DEFAULT_STORAGE_SYSTEM_COST_YUAN_PER_WH,
        "power_kw": round(power, 2),
        "duration_hours": round(duration or 2, 2),
        "capacity_kwh": round(capacity or power * (duration or 2), 2),
        "dod": 0.95,
        "system_efficiency": 0.88,
        "annual_operating_days": 320,
        "availability": 0.95,
        "annual_degradation": 0.02,
        "project_years": 15,
        "battery_unit_cost": DEFAULT_NA_CELL_COST_YUAN_PER_WH,
        "pcs_ems_bms_unit_cost": DEFAULT_PCS_EMS_BMS_COST_YUAN_PER_WH,
        "grid_connection_cost_per_kw": DEFAULT_GRID_CONNECTION_COST_PER_KW,
        "civil_fire_cost_per_kw": DEFAULT_CIVIL_FIRE_COST_PER_KW,
        "design_supervision_rate": DEFAULT_INCLUDED_COST_RATE,
        "contingency_rate": DEFAULT_INCLUDED_COST_RATE,
        "construction_months": 6,
        "construction_interest_rate": DEFAULT_INCLUDED_COST_RATE,
        "valley_charge_price": 0.275,
        "flat_charge_price": 0.612,
        "discharge_price": 0.919,
        "valley_peak_cycles": 1.4,
        "flat_peak_cycles": 0,
        "daily_cycles": 1.4,
        "price_escalation": 0,
        "demand_revenue": 0,
        "demand_revenue_per_kw_year": 120,
        "ancillary_revenue": 0,
        "ancillary_revenue_per_mw_year": 0,
        "other_revenue": 0,
        "variable_om_cost_per_kwh": 0.015,
        "om_cost_rate": 0.015,
        "insurance_land_mgmt_rate": 0.005,
        "om_escalation": 0,
        "tax_rate": 0.25,
        "depreciation_years": 10,
        "fixed_asset_residual_rate": 0.05,
        "discount_rate": 0.08,
        "terminal_residual_rate": 0.05,
        "residual_rate": 0.05,
        "battery_replacement_year": 0,
        "battery_replacement_cost_rate": 0.45,
        "enable_customer_share": True,
        "customer_share_before_payback": 0.2,
        "customer_share_after_payback": 0.5,
        "enable_loan": True,
        "loan_ratio": 0.7,
        "loan_interest_rate": 0.042,
        "loan_years": 8,
    }


def _as_rate(value: Any, default: float = 0.0) -> float:
    raw = _safe_float(value, 8, default)
    if raw is None:
        return default
    return raw / 100 if abs(raw) > 1.5 else raw


def _bool_param(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return value.strip().lower() in {"1", "true", "yes", "y", "on", "是", "启用"}
    return bool(value)


def _merge_revenue_params(agent: StorageAgent, req: dict | None) -> dict:
    params = _default_revenue_params(agent)
    req = dict(req or {})
    aliases = {
        "valley_price": "valley_charge_price",
        "flat_price": "flat_charge_price",
        "peak_price": "discharge_price",
    }
    if "daily_cycles" in req and "valley_peak_cycles" not in req and "flat_peak_cycles" not in req:
        req["valley_peak_cycles"] = req.get("daily_cycles")
        req["flat_peak_cycles"] = 0
    for old_key, new_key in aliases.items():
        if old_key in req and new_key not in req:
            req[new_key] = req[old_key]
    for key, value in req.items():
        if key == "user_id":
            continue
        if key in params:
            params[key] = value
    for key in ("enable_customer_share", "enable_loan"):
        params[key] = _bool_param(params.get(key))
    for key in (
        "dod", "system_efficiency", "availability", "annual_degradation",
        "design_supervision_rate", "contingency_rate", "construction_interest_rate",
        "price_escalation", "om_cost_rate", "insurance_land_mgmt_rate", "om_escalation",
        "tax_rate", "fixed_asset_residual_rate", "discount_rate", "terminal_residual_rate",
        "residual_rate", "battery_replacement_cost_rate", "customer_share_before_payback",
        "customer_share_after_payback", "loan_ratio", "loan_interest_rate",
    ):
        params[key] = _as_rate(params.get(key), 0.0)
    for key in (
        "power_kw", "duration_hours", "capacity_kwh", "annual_operating_days",
        "project_years", "battery_unit_cost", "pcs_ems_bms_unit_cost",
        "grid_connection_cost_per_kw", "civil_fire_cost_per_kw", "construction_months",
        "valley_charge_price", "flat_charge_price", "discharge_price",
        "valley_peak_cycles", "flat_peak_cycles", "demand_revenue",
        "demand_revenue_per_kw_year", "ancillary_revenue",
        "ancillary_revenue_per_mw_year", "other_revenue", "variable_om_cost_per_kwh",
        "depreciation_years", "battery_replacement_year", "loan_years",
        "cell_cost_yuan_per_wh", "system_unit_cost_yuan_per_wh",
    ):
        params[key] = _safe_float(params.get(key), 6) or 0
    if params["battery_unit_cost"] > 20:
        params["battery_unit_cost"] = params["battery_unit_cost"] / 1000
    if params["pcs_ems_bms_unit_cost"] > 20:
        params["pcs_ems_bms_unit_cost"] = params["pcs_ems_bms_unit_cost"] / 1000
    if params["grid_connection_cost_per_kw"] > 20000:
        params["grid_connection_cost_per_kw"] = params["grid_connection_cost_per_kw"] / max(params["power_kw"], 1)
    if params["civil_fire_cost_per_kw"] > 20000:
        params["civil_fire_cost_per_kw"] = params["civil_fire_cost_per_kw"] / max(params["power_kw"], 1)
    params["power_kw"] = max(0.0, params["power_kw"])
    params["duration_hours"] = max(0.0, params["duration_hours"])
    if not params.get("capacity_kwh"):
        params["capacity_kwh"] = params["power_kw"] * params["duration_hours"]
    params["capacity_kwh"] = max(0.0, _safe_float(params.get("capacity_kwh"), 4) or 0)
    if "pcs_ems_bms_cost_per_kw" in req and "pcs_ems_bms_unit_cost" not in req:
        old_pcs = _safe_float(req.get("pcs_ems_bms_cost_per_kw"), 6) or 0
        if params["capacity_kwh"] > 0:
            params["pcs_ems_bms_unit_cost"] = old_pcs * params["power_kw"] / (params["capacity_kwh"] * 1000)
    if "grid_connection_cost" in req and "grid_connection_cost_per_kw" not in req:
        old_grid = _safe_float(req.get("grid_connection_cost"), 6) or 0
        params["grid_connection_cost_per_kw"] = old_grid / max(params["power_kw"], 1)
    if "civil_fire_cost" in req and "civil_fire_cost_per_kw" not in req:
        old_civil = _safe_float(req.get("civil_fire_cost"), 6) or 0
        params["civil_fire_cost_per_kw"] = old_civil / max(params["power_kw"], 1)
    params["daily_cycles"] = max(0.0, params["valley_peak_cycles"]) + max(0.0, params["flat_peak_cycles"])
    params["valley_price"] = params["valley_charge_price"]
    params["flat_price"] = params["flat_charge_price"]
    params["peak_price"] = params["discharge_price"]
    params["model_version"] = REVENUE_MODEL_VERSION
    params["model_source"] = REVENUE_TEMPLATE_PATH
    return params


def _loan_schedule(loan_amount: float, annual_rate: float, loan_years: int, total_years: int) -> list[dict]:
    rows: list[dict] = []
    balance = max(0.0, loan_amount)
    principal_each_year = loan_amount / loan_years if loan_amount > 0 and loan_years > 0 else 0.0
    for year in range(1, total_years + 1):
        beginning = balance
        active = beginning > 0 and year <= loan_years
        interest = beginning * annual_rate if active else 0.0
        principal = min(beginning, principal_each_year) if active else 0.0
        balance = max(0.0, beginning - principal)
        rows.append({
            "year": year,
            "status": "未结清" if beginning > 0 else ("无贷款" if loan_amount <= 0 else "已结清"),
            "beginning_balance": beginning,
            "interest": interest,
            "principal": principal,
            "debt_service": interest + principal,
            "ending_balance": balance,
        })
    return rows


def _irr(cash_flows: list[float], max_iter: int = 1000, tol: float = 1e-7) -> float | None:
    if not cash_flows or not any(v < 0 for v in cash_flows) or not any(v > 0 for v in cash_flows):
        return None
    rate = 0.1
    for _ in range(max_iter):
        if rate <= -0.99:
            rate = -0.9
        npv = sum(cf / (1 + rate) ** i for i, cf in enumerate(cash_flows))
        dnpv = sum(-i * cf / (1 + rate) ** (i + 1) for i, cf in enumerate(cash_flows))
        if abs(dnpv) < 1e-12:
            break
        new_rate = rate - npv / dnpv
        if not math.isfinite(new_rate):
            break
        if abs(new_rate - rate) < tol:
            return new_rate
        rate = new_rate
    return rate if math.isfinite(rate) else None


def _compute_revenue_model(agent: StorageAgent, req: dict | None = None) -> dict:
    if _normalize_bill_df(agent.state.electricity_df).empty:
        _load_bills(agent)
    req = dict(req or {})
    sync_to_capacity = _bool_param(req.pop("_sync_to_capacity", False))
    if sync_to_capacity:
        req.update({
            "battery_unit_cost": DEFAULT_NA_CELL_COST_YUAN_PER_WH,
            "pcs_ems_bms_unit_cost": DEFAULT_PCS_EMS_BMS_COST_YUAN_PER_WH,
            "grid_connection_cost_per_kw": DEFAULT_GRID_CONNECTION_COST_PER_KW,
            "civil_fire_cost_per_kw": DEFAULT_CIVIL_FIRE_COST_PER_KW,
            "cell_cost_yuan_per_wh": DEFAULT_NA_CELL_COST_YUAN_PER_WH,
            "system_unit_cost_yuan_per_wh": DEFAULT_STORAGE_SYSTEM_COST_YUAN_PER_WH,
            "cost_basis": DEFAULT_STORAGE_COST_BASIS,
            "design_supervision_rate": DEFAULT_INCLUDED_COST_RATE,
            "contingency_rate": DEFAULT_INCLUDED_COST_RATE,
            "construction_interest_rate": DEFAULT_INCLUDED_COST_RATE,
        })
    params = _merge_revenue_params(agent, req)
    capacity = float(params["capacity_kwh"])
    power = float(params["power_kw"])
    duration = float(params["duration_hours"]) or (capacity / power if power else 0)
    dod = min(1.0, max(0.0, float(params["dod"])))
    efficiency = min(1.0, max(0.01, float(params["system_efficiency"])))
    days = max(0, int(float(params["annual_operating_days"])))
    years = max(1, int(float(params["project_years"])))
    availability = min(1.0, max(0.0, float(params["availability"])))
    degradation = min(0.99, max(0.0, float(params["annual_degradation"])))
    valley_cycles = max(0.0, float(params["valley_peak_cycles"]))
    flat_cycles = max(0.0, float(params["flat_peak_cycles"]))
    cycles = valley_cycles + flat_cycles
    valley_charge_price = max(0.0, float(params["valley_charge_price"]))
    flat_charge_price = max(0.0, float(params["flat_charge_price"]))
    discharge_price = max(0.0, float(params["discharge_price"]))
    charge_price = (
        (valley_charge_price * valley_cycles + flat_charge_price * flat_cycles) / cycles
        if cycles > 0 else valley_charge_price
    )
    spread = max(0.0, discharge_price - charge_price)

    battery_cost = capacity * 1000 * float(params["battery_unit_cost"])
    pcs_cost = capacity * 1000 * float(params["pcs_ems_bms_unit_cost"])
    grid_cost = power * float(params["grid_connection_cost_per_kw"])
    civil_cost = power * float(params["civil_fire_cost_per_kw"])
    direct_cost = battery_cost + pcs_cost + grid_cost + civil_cost
    target_direct_cost = capacity * 1000 * float(params.get("system_unit_cost_yuan_per_wh") or 0)
    if target_direct_cost > 0:
        battery_cost = min(battery_cost, target_direct_cost)
        remaining_direct_cost = max(0.0, target_direct_cost - battery_cost)
        pcs_cost = min(pcs_cost, remaining_direct_cost)
        remaining_direct_cost = max(0.0, remaining_direct_cost - pcs_cost)
        grid_cost = min(grid_cost, remaining_direct_cost)
        remaining_direct_cost = max(0.0, remaining_direct_cost - grid_cost)
        civil_cost = remaining_direct_cost
        direct_cost = target_direct_cost
    design_cost = direct_cost * float(params["design_supervision_rate"])
    after_design = direct_cost + design_cost
    contingency = after_design * float(params["contingency_rate"])
    after_contingency = after_design + contingency
    construction_interest = after_contingency * float(params["construction_interest_rate"]) * float(params["construction_months"]) / 12 / 2
    total_investment = after_contingency + construction_interest

    base_discharge = capacity * dod * days * cycles * availability
    base_charge = base_discharge / efficiency
    base_demand_revenue = (
        float(params["demand_revenue"])
        if float(params["demand_revenue"]) > 0
        else power * float(params["demand_revenue_per_kw_year"])
    )
    base_ancillary_revenue = (
        float(params["ancillary_revenue"])
        if float(params["ancillary_revenue"]) > 0
        else power / 1000 * float(params["ancillary_revenue_per_mw_year"])
    )
    other_revenue = float(params["other_revenue"])
    price_escalation = float(params["price_escalation"])
    om_escalation = float(params["om_escalation"])
    tax_rate = float(params["tax_rate"])
    depreciation_years = max(1, int(float(params.get("depreciation_years") or years)))
    depreciation_residual_value = total_investment * float(params["fixed_asset_residual_rate"])
    terminal_residual_value = total_investment * float(params["terminal_residual_rate"])
    annual_depreciation = max(0.0, (total_investment - depreciation_residual_value) / depreciation_years)
    before_share = float(params["customer_share_before_payback"]) if params.get("enable_customer_share") else 0
    after_share = float(params["customer_share_after_payback"]) if params.get("enable_customer_share") else 0

    loan_enabled = bool(params.get("enable_loan"))
    loan_ratio = min(1.0, max(0.0, float(params["loan_ratio"])))
    loan_years = max(0, int(float(params["loan_years"])))
    loan_rate = max(0.0, float(params["loan_interest_rate"]))
    loan_amount = total_investment * loan_ratio if loan_enabled else 0.0
    equity_amount = total_investment - loan_amount
    current_schedule = _loan_schedule(loan_amount, loan_rate, loan_years, years)
    sim_loan_amount = total_investment * loan_ratio
    sim_schedule = _loan_schedule(sim_loan_amount, loan_rate, loan_years, years)

    base_rows = []
    project_flows = [-total_investment]
    discounted_project_flows = [-total_investment]
    cumulative_project = -total_investment
    cumulative_discounted = -total_investment
    discount_rate = max(0.0, float(params["discount_rate"]))
    simple_payback_year = None
    dynamic_payback_year = None
    share_cumulative_investor = -equity_amount
    for year in range(1, years + 1):
        debt_for_share = current_schedule[year - 1] if loan_enabled and year - 1 < len(current_schedule) else {
            "interest": 0.0,
            "principal": 0.0,
        }
        capacity_factor = availability * (1 - degradation) ** (year - 1)
        discharge = capacity * dod * days * cycles * capacity_factor
        charge = discharge / efficiency
        revenue_factor = (1 + price_escalation) ** (year - 1)
        cost_factor = (1 + om_escalation) ** (year - 1)
        sale_revenue = discharge * discharge_price * revenue_factor
        demand_revenue = base_demand_revenue * revenue_factor
        ancillary_revenue = base_ancillary_revenue * revenue_factor
        fixed_other_revenue = other_revenue * revenue_factor
        gross = sale_revenue + demand_revenue + ancillary_revenue + fixed_other_revenue
        charging = charge * charge_price * revenue_factor
        fixed_om = total_investment * float(params["om_cost_rate"]) * cost_factor
        variable_om = discharge * float(params["variable_om_cost_per_kwh"]) * cost_factor
        insurance_cost = total_investment * float(params["insurance_land_mgmt_rate"]) * cost_factor
        operating_cost = charging + fixed_om + variable_om + insurance_cost
        ebitda_before_share = gross - operating_cost
        share_ratio = after_share if share_cumulative_investor >= 0 else before_share
        customer_share_original = max(0.0, ebitda_before_share) * share_ratio
        investor_ebitda = ebitda_before_share - customer_share_original
        depreciation = annual_depreciation if year <= depreciation_years else 0.0
        taxable_income = max(0.0, investor_ebitda - depreciation)
        tax = taxable_income * tax_rate
        replacement_cash = 0.0
        if int(float(params["battery_replacement_year"] or 0)) == year:
            replacement_cash = -battery_cost * float(params["battery_replacement_cost_rate"])
        net_operating_cash = investor_ebitda - tax
        residual_cash = terminal_residual_value if year == years else 0.0
        investor_cash_for_share = (
            net_operating_cash + residual_cash + replacement_cash
            - debt_for_share["interest"] - debt_for_share["principal"]
        )
        project_net = replacement_cash + net_operating_cash + residual_cash
        cumulative_project += project_net
        discounted = project_net / (1 + discount_rate) ** year if discount_rate > -1 else project_net
        cumulative_discounted += discounted
        if simple_payback_year is None and cumulative_project >= 0:
            simple_payback_year = year
        if dynamic_payback_year is None and cumulative_discounted >= 0:
            dynamic_payback_year = year
        row = {
            "year": year,
            "capacity_factor": capacity_factor,
            "discharge_kwh": discharge,
            "charge_kwh": charge,
            "sale_revenue_yuan": sale_revenue,
            "arbitrage_margin_yuan": sale_revenue - charging,
            "demand_revenue_yuan": demand_revenue,
            "ancillary_revenue_yuan": ancillary_revenue,
            "other_revenue_yuan": fixed_other_revenue,
            "gross_revenue_yuan": gross,
            "charge_cost_yuan": charging,
            "fixed_om_yuan": fixed_om,
            "variable_om_yuan": variable_om,
            "insurance_cost_yuan": insurance_cost,
            "operating_cost_yuan": operating_cost,
            "ebitda_before_share_yuan": ebitda_before_share,
            "share_ratio": share_ratio,
            "customer_share_original_yuan": customer_share_original,
            "investor_ebitda_yuan": investor_ebitda,
            "depreciation_yuan": depreciation,
            "tax_yuan": tax,
            "replacement_cash_yuan": replacement_cash,
            "net_operating_cash_yuan": net_operating_cash,
            "residual_value_yuan": residual_cash,
            "project_net_cash_flow_yuan": project_net,
            "discounted_project_cash_flow_yuan": discounted,
            "cumulative_project_yuan": cumulative_project,
            "cumulative_discounted_project_yuan": cumulative_discounted,
        }
        base_rows.append(row)
        project_flows.append(project_net)
        discounted_project_flows.append(discounted)
        share_cumulative_investor += investor_cash_for_share

    project_irr = _irr(project_flows)
    project_npv = sum(project_flows[i] / (1 + discount_rate) ** i for i in range(len(project_flows)))

    investor_flows = [-equity_amount]
    sim_loan_flows = [-(total_investment - sim_loan_amount)]
    no_loan_flows = [-total_investment]
    cash_flow = []
    customer_yearly = []
    investor_yearly = []
    cumulative_investor = investor_flows[0]
    cumulative_customer = 0.0
    cumulative_investor_adjusted = 0.0
    cumulative_customer_original = 0.0
    cumulative_sim_loan = sim_loan_flows[0]
    cumulative_no_loan = no_loan_flows[0]

    for row, current_debt, sim_debt in zip(base_rows, current_schedule, sim_schedule):
        investor_cash = (
            row["net_operating_cash_yuan"]
            + row["residual_value_yuan"]
            + row["replacement_cash_yuan"]
            - current_debt["interest"]
            - current_debt["principal"]
        )
        cumulative_investor += investor_cash
        investor_flows.append(investor_cash)

        sim_loan_cash = row["project_net_cash_flow_yuan"] - sim_debt["interest"] - sim_debt["principal"]
        no_loan_cash = row["project_net_cash_flow_yuan"]
        cumulative_sim_loan += sim_loan_cash
        cumulative_no_loan += no_loan_cash
        sim_loan_flows.append(sim_loan_cash)
        no_loan_flows.append(no_loan_cash)

        loan_cost = current_debt["debt_service"] if loan_enabled else 0.0
        distributable = max(row["ebitda_before_share_yuan"] - loan_cost, 0.0)
        customer_adjusted = distributable * row["share_ratio"] if params.get("enable_customer_share") else 0.0
        investor_adjusted = max(distributable - customer_adjusted, 0.0)
        cumulative_customer += customer_adjusted
        cumulative_investor_adjusted += investor_adjusted
        cumulative_customer_original += row["customer_share_original_yuan"]
        customer_yearly.append({
            "year": row["year"],
            "loan_status": current_debt["status"] if loan_enabled else "无贷款",
            "loan_cost_yuan": _safe_float(loan_cost, 2),
            "share_ratio_percent": _safe_float(row["share_ratio"] * 100, 2),
            "distributable_yuan": _safe_float(distributable, 2),
            "customer_income_yuan": _safe_float(customer_adjusted, 2),
            "cumulative_customer_yuan": _safe_float(cumulative_customer, 2),
            "original_customer_income_yuan": _safe_float(row["customer_share_original_yuan"], 2),
            "loan_impact_yuan": _safe_float(row["customer_share_original_yuan"] - customer_adjusted, 2),
        })
        investor_yearly.append({
            "year": row["year"],
            "loan_status": current_debt["status"] if loan_enabled else "无贷款",
            "beginning_loan_balance_yuan": _safe_float(current_debt["beginning_balance"] if loan_enabled else 0, 2),
            "loan_cost_yuan": _safe_float(loan_cost, 2),
            "distributable_yuan": _safe_float(distributable, 2),
            "investor_income_yuan": _safe_float(investor_adjusted, 2),
            "cumulative_investor_income_yuan": _safe_float(cumulative_investor_adjusted, 2),
            "original_investor_income_yuan": _safe_float(row["investor_ebitda_yuan"], 2),
            "loan_impact_yuan": _safe_float(row["investor_ebitda_yuan"] - investor_adjusted, 2),
        })
        cash_flow.append({
            "year": row["year"],
            "capacity_factor_percent": _safe_float(row["capacity_factor"] * 100, 2),
            "discharge_kwh": _safe_float(row["discharge_kwh"], 2),
            "charge_kwh": _safe_float(row["charge_kwh"], 2),
            "sale_revenue_yuan": _safe_float(row["sale_revenue_yuan"], 2),
            "energy_revenue_yuan": _safe_float(row["sale_revenue_yuan"], 2),
            "arbitrage_margin_yuan": _safe_float(row["arbitrage_margin_yuan"], 2),
            "demand_revenue_yuan": _safe_float(row["demand_revenue_yuan"], 2),
            "ancillary_revenue_yuan": _safe_float(row["ancillary_revenue_yuan"], 2),
            "other_revenue_yuan": _safe_float(row["other_revenue_yuan"], 2),
            "gross_revenue_yuan": _safe_float(row["gross_revenue_yuan"], 2),
            "charge_cost_yuan": _safe_float(row["charge_cost_yuan"], 2),
            "om_cost_yuan": _safe_float(row["fixed_om_yuan"] + row["variable_om_yuan"] + row["insurance_cost_yuan"], 2),
            "customer_share_yuan": _safe_float(row["customer_share_original_yuan"], 2),
            "investor_ebitda_yuan": _safe_float(row["investor_ebitda_yuan"], 2),
            "depreciation_yuan": _safe_float(row["depreciation_yuan"], 2),
            "tax_yuan": _safe_float(row["tax_yuan"], 2),
            "replacement_cash_yuan": _safe_float(row["replacement_cash_yuan"], 2),
            "residual_value_yuan": _safe_float(row["residual_value_yuan"], 2),
            "loan_interest_yuan": _safe_float(current_debt["interest"] if loan_enabled else 0, 2),
            "loan_principal_yuan": _safe_float(current_debt["principal"] if loan_enabled else 0, 2),
            "debt_service_yuan": _safe_float(current_debt["debt_service"] if loan_enabled else 0, 2),
            "investor_cash_flow_yuan": _safe_float(investor_cash, 2),
            "project_net_cash_flow_yuan": _safe_float(row["project_net_cash_flow_yuan"], 2),
            "cumulative_project_yuan": _safe_float(row["cumulative_project_yuan"], 2),
            "cumulative_investor_yuan": _safe_float(cumulative_investor, 2),
        })

    investor_irr = _irr(investor_flows)
    sim_loan_irr = _irr(sim_loan_flows)
    no_loan_irr = _irr(no_loan_flows)
    first = base_rows[0] if base_rows else {}
    first_investor_ebitda = float(first.get("investor_ebitda_yuan") or 0)
    static_payback = total_investment / first_investor_ebitda if first_investor_ebitda > 0 else None
    customer_total = sum(r["customer_income_yuan"] for r in customer_yearly)
    investor_total = sum(r["investor_income_yuan"] for r in investor_yearly)
    total_interest = sum(r["interest"] for r in current_schedule) if loan_enabled else 0.0
    total_principal = sum(r["principal"] for r in current_schedule) if loan_enabled else 0.0
    sim_total_interest = sum(r["interest"] for r in sim_schedule)
    sim_total_payment = sum(r["debt_service"] for r in sim_schedule)
    loan_first_year = sim_loan_flows[1] if len(sim_loan_flows) > 1 else 0.0
    no_loan_first_year = no_loan_flows[1] if len(no_loan_flows) > 1 else 0.0
    loan_lifetime_total = sum(sim_loan_flows[1:])
    no_loan_lifetime_total = sum(no_loan_flows[1:])
    loan_lifetime_net = sum(sim_loan_flows)
    no_loan_lifetime_net = sum(no_loan_flows)
    irr_delta = (sim_loan_irr or 0) - (no_loan_irr or 0)
    lifetime_delta = loan_lifetime_total - no_loan_lifetime_total
    net_delta = loan_lifetime_net - no_loan_lifetime_net
    if not loan_enabled:
        recommended_ratio = "当前方案未使用贷款，建议以无贷款结果为准"
    elif irr_delta <= 0:
        recommended_ratio = "建议0%-30%或不贷款"
    elif net_delta < 0:
        recommended_ratio = "建议30%-50%，兼顾IRR与净现金流"
    elif lifetime_delta < 0:
        recommended_ratio = "建议40%-60%，兼顾IRR与总收益"
    else:
        recommended_ratio = "建议50%-70%，杠杆效果较优"
    loan_conclusion = (
        "当前方案未使用贷款，因此贷款不会对当前收益产生实际影响；上方模拟有贷款/无贷款结果仅用于融资方案对比。"
        if not loan_enabled else
        ("贷款能够提高股东IRR；" if irr_delta > 0 else "贷款未提高股东IRR；")
        + ("会压缩资方生命周期总收益；" if lifetime_delta < 0 else "不会压缩资方生命周期总收益；")
        + ("同时会降低资方生命周期净现金流。" if net_delta < 0 else "且不会降低资方生命周期净现金流。")
    )

    sensitivity = []
    first_sale = float(first.get("sale_revenue_yuan") or 0)
    first_charging = float(first.get("charge_cost_yuan") or 0)
    first_margin = first_sale - first_charging
    first_ebitda_before = float(first.get("ebitda_before_share_yuan") or 0)
    for factor_name, values in {
        "电价变化": [-0.2, -0.1, 0, 0.1, 0.2],
        "投资成本变化": [-0.2, -0.1, 0, 0.1, 0.2],
        "循环次数变化": [-0.3, -0.15, 0, 0.15, 0.3],
        "分成比例变化": [-0.1, -0.05, 0, 0.05, 0.1],
        "贷款比例变化": [-0.2, -0.1, 0, 0.1, 0.2],
    }.items():
        for delta in values:
            adj_net = first_investor_ebitda
            adj_investment = total_investment
            if factor_name == "电价变化":
                adj_net = first_investor_ebitda + first_margin * delta
            elif factor_name == "投资成本变化":
                adj_investment = total_investment * (1 + delta)
                adj_net = first_investor_ebitda - total_investment * (
                    float(params["om_cost_rate"]) + float(params["insurance_land_mgmt_rate"])
                ) * delta
            elif factor_name == "循环次数变化":
                adj_net = first_investor_ebitda + first_margin * delta
            elif factor_name == "分成比例变化":
                adj_share = max(0.0, before_share + delta)
                adj_net = first_ebitda_before - max(first_ebitda_before, 0) * adj_share
            elif factor_name == "贷款比例变化":
                adj_net = first_investor_ebitda - (sim_schedule[0]["debt_service"] if sim_schedule else 0) * delta
            sensitivity.append({
                "factor": factor_name,
                "change_percent": _safe_float(delta * 100, 0),
                "annual_net_yuan": _safe_float(adj_net, 2),
                "payback_years": _safe_float(adj_investment / adj_net, 2, None) if adj_net > 0 else None,
            })

    cycle_sensitivity = []
    for annual_cycles in [180, 240, 300, 365, 400, 450, 500, 600]:
        discharge = capacity * dod * availability * annual_cycles
        charge = discharge / efficiency
        sale = discharge * discharge_price
        charging = charge * charge_price
        margin = sale - charging
        fixed_om = total_investment * float(params["om_cost_rate"])
        variable_om = discharge * float(params["variable_om_cost_per_kwh"])
        insurance_cost = total_investment * float(params["insurance_land_mgmt_rate"])
        ebitda = margin + base_demand_revenue + base_ancillary_revenue + other_revenue - fixed_om - variable_om - insurance_cost
        investor_income = ebitda - max(ebitda, 0) * before_share
        cycle_sensitivity.append({
            "annual_cycles": annual_cycles,
            "discharge_kwh": _safe_float(discharge, 2),
            "arbitrage_revenue_yuan": _safe_float(margin, 2),
            "demand_revenue_yuan": _safe_float(base_demand_revenue, 2),
            "annual_net_yuan": _safe_float(investor_income, 2),
            "payback_years": _safe_float(total_investment / investor_income, 2, None) if investor_income > 0 else None,
        })

    payback_matrix = {
        "spreads": [0.3, 0.4, 0.5, 0.6, 0.7, 0.8],
        "rows": [],
    }
    for unit_cost in [0.65, 0.75, 0.85, 0.95, 1.05, 1.15]:
        matrix_values = []
        matrix_direct = capacity * 1000 * unit_cost + power * (
            float(params["grid_connection_cost_per_kw"]) + float(params["civil_fire_cost_per_kw"])
        )
        matrix_capex = matrix_direct * (1 + float(params["design_supervision_rate"])) * (1 + float(params["contingency_rate"]))
        matrix_capex *= 1 + float(params["construction_interest_rate"]) * float(params["construction_months"]) / 12 / 2
        for matrix_spread in payback_matrix["spreads"]:
            matrix_sale_price = charge_price + matrix_spread
            matrix_gross_margin = base_discharge * matrix_sale_price - base_charge * charge_price
            matrix_ebitda = (
                matrix_gross_margin + base_demand_revenue + base_ancillary_revenue + other_revenue
                - base_discharge * float(params["variable_om_cost_per_kwh"])
                - matrix_capex * (float(params["om_cost_rate"]) + float(params["insurance_land_mgmt_rate"]))
            )
            matrix_investor = matrix_ebitda - max(matrix_ebitda, 0) * before_share
            matrix_values.append(_safe_float(matrix_capex / matrix_investor, 2, None) if matrix_investor > 0 else None)
        payback_matrix["rows"].append({"unit_cost_yuan_per_wh": unit_cost, "values": matrix_values})

    share_sensitivity = []
    for share in [0, 0.1, 0.2, 0.3, 0.4, 0.5]:
        income = first_ebitda_before - max(first_ebitda_before, 0) * share
        share_sensitivity.append({
            "share_ratio_percent": _safe_float(share * 100, 0),
            "payback_years": _safe_float(total_investment / income, 2, None) if income > 0 else None,
        })

    result = {
        "ok": True,
        "model_version": REVENUE_MODEL_VERSION,
        "model_source": REVENUE_TEMPLATE_PATH,
        "params_source": "capacity_sync" if sync_to_capacity else "manual",
        "params": params,
        "summary": {
            "power_kw": _safe_float(power, 2),
            "duration_hours": _safe_float(duration, 2),
            "capacity_kwh": _safe_float(capacity, 2),
            "total_investment_yuan": _safe_float(total_investment, 2),
            "unit_investment_yuan_per_wh": _safe_float(total_investment / (capacity * 1000), 4, None) if capacity > 0 else None,
            "charge_price_yuan_per_kwh": _safe_float(charge_price, 4),
            "discharge_price_yuan_per_kwh": _safe_float(discharge_price, 4),
            "spread_yuan_per_kwh": _safe_float(spread, 4),
            "daily_cycles": _safe_float(cycles, 3),
            "first_year_discharge_kwh": _safe_float(first.get("discharge_kwh"), 2),
            "first_year_charge_kwh": _safe_float(first.get("charge_kwh"), 2),
            "sale_revenue_yuan": _safe_float(first.get("sale_revenue_yuan"), 2),
            "charge_cost_yuan": _safe_float(first.get("charge_cost_yuan"), 2),
            "arbitrage_revenue_yuan": _safe_float(first_margin, 2),
            "demand_revenue_yuan": _safe_float(first.get("demand_revenue_yuan"), 2),
            "first_year_gross_revenue_yuan": _safe_float(first.get("gross_revenue_yuan"), 2),
            "first_year_operating_cost_yuan": _safe_float(first.get("operating_cost_yuan"), 2),
            "first_year_ebitda_yuan": _safe_float(first_ebitda_before, 2),
            "first_year_investor_ebitda_yuan": _safe_float(first_investor_ebitda, 2),
            "first_year_net_yuan": _safe_float(cash_flow[0]["investor_cash_flow_yuan"] if cash_flow else 0, 2),
            "annual_depreciation_yuan": _safe_float(annual_depreciation, 2),
            "residual_value_yuan": _safe_float(terminal_residual_value, 2),
            "customer_first_year_yuan": _safe_float(customer_yearly[0]["customer_income_yuan"] if customer_yearly else 0, 2),
            "customer_first_year_original_yuan": _safe_float(first.get("customer_share_original_yuan"), 2),
            "investor_first_year_yuan": _safe_float(investor_yearly[0]["investor_income_yuan"] if investor_yearly else 0, 2),
            "investor_first_year_original_yuan": _safe_float(first_investor_ebitda, 2),
            "static_payback_years": _safe_float(static_payback, 2, None),
            "simple_payback_year": simple_payback_year,
            "dynamic_payback_year": dynamic_payback_year,
            "project_npv_yuan": _safe_float(project_npv, 2),
            "project_irr_percent": _safe_float((project_irr or 0) * 100, 2, None),
            "investor_irr_percent": _safe_float((investor_irr or 0) * 100, 2, None),
            "customer_lifetime_yuan": _safe_float(customer_total, 2),
            "investor_lifetime_yuan": _safe_float(investor_total, 2),
        },
        "loan": {
            "enabled": loan_enabled,
            "loan_amount_yuan": _safe_float(loan_amount, 2),
            "equity_amount_yuan": _safe_float(equity_amount, 2),
            "annual_debt_service_yuan": _safe_float(current_schedule[0]["debt_service"] if current_schedule and loan_enabled else 0, 2),
            "total_interest_yuan": _safe_float(total_interest, 2),
            "total_payment_yuan": _safe_float(total_principal + total_interest, 2),
            "recommended_ratio_range": recommended_ratio,
            "conclusion": loan_conclusion,
        },
        "loan_impact": {
            "loan_first_year_cash_flow_yuan": _safe_float(loan_first_year, 2),
            "no_loan_first_year_cash_flow_yuan": _safe_float(no_loan_first_year, 2),
            "first_year_delta_yuan": _safe_float(loan_first_year - no_loan_first_year, 2),
            "loan_lifetime_total_yuan": _safe_float(loan_lifetime_total, 2),
            "no_loan_lifetime_total_yuan": _safe_float(no_loan_lifetime_total, 2),
            "lifetime_delta_yuan": _safe_float(lifetime_delta, 2),
            "loan_lifetime_net_cash_flow_yuan": _safe_float(loan_lifetime_net, 2),
            "no_loan_lifetime_net_cash_flow_yuan": _safe_float(no_loan_lifetime_net, 2),
            "net_delta_yuan": _safe_float(net_delta, 2),
            "loan_shareholder_irr_percent": _safe_float((sim_loan_irr or 0) * 100, 2, None),
            "no_loan_shareholder_irr_percent": _safe_float((no_loan_irr or 0) * 100, 2, None),
            "irr_delta_percent": _safe_float(irr_delta * 100, 2, None),
            "total_interest_yuan": _safe_float(sim_total_interest, 2),
            "total_payment_yuan": _safe_float(sim_total_payment, 2),
            "recommended_ratio_range": recommended_ratio,
            "conclusion": loan_conclusion,
        },
        "cash_flow": cash_flow,
        "customer_yearly": customer_yearly,
        "investor_yearly": investor_yearly,
        "sensitivity": sensitivity,
        "cycle_sensitivity": cycle_sensitivity,
        "payback_matrix": payback_matrix,
        "share_sensitivity": share_sensitivity,
        "cost_breakdown": [
            {"item": "电池系统", "amount_yuan": _safe_float(battery_cost, 2)},
            {"item": "PCS/EMS/BMS", "amount_yuan": _safe_float(pcs_cost, 2)},
            {"item": "变配电及并网", "amount_yuan": _safe_float(grid_cost, 2)},
            {"item": "土建/消防/安装", "amount_yuan": _safe_float(civil_cost, 2)},
            {"item": "设计监理及其他", "amount_yuan": _safe_float(design_cost, 2)},
            {"item": "预备费", "amount_yuan": _safe_float(contingency, 2)},
            {"item": "建设期资金成本", "amount_yuan": _safe_float(construction_interest, 2)},
            {"item": "总投资", "amount_yuan": _safe_float(total_investment, 2)},
        ],
    }
    agent.state.revenue_model = result
    return _persist_revenue(agent, result)


def _md_value(value: Any, digits: int = 2) -> str:
    if value is None:
        return "-"
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, (int, float)):
        if not math.isfinite(float(value)):
            return "-"
        text = f"{float(value):,.{digits}f}"
        if "." in text:
            text = text.rstrip("0").rstrip(".")
        return text
    text = str(value)
    return text.replace("|", "\\|").replace("\n", " ")


def _md_table(rows: list[dict], columns: list[tuple[str, str]], *, limit: int | None = None) -> str:
    data = rows[:limit] if limit else rows
    if not data:
        return "暂无数据\n"
    header = "| " + " | ".join(label for _, label in columns) + " |"
    divider = "| " + " | ".join("---" for _ in columns) + " |"
    body = [
        "| " + " | ".join(_md_value(row.get(key), 2) for key, _ in columns) + " |"
        for row in data
    ]
    return "\n".join([header, divider, *body]) + "\n"


def _report_num(value: Any, digits: int = 2) -> str:
    return _md_value(value, digits)


def _report_money(value: Any, digits: int = 2) -> str:
    return f"{_md_value(value, digits)} 元"


def _report_percent(value: Any, digits: int = 2) -> str:
    return f"{_md_value(value, digits)}%"


def _report_rows(rows: list[dict], columns: list[tuple[str, str]], *, limit: int | None = None) -> list[list[str]]:
    data = rows[:limit] if limit else rows
    return [[_md_value(row.get(key), 2) for key, _ in columns] for row in data]


def _bill_records_for_report(agent: StorageAgent) -> list[dict]:
    bills = _load_bills(agent) or {}
    rows = bills.get("records") or []
    if rows:
        return rows
    df = _normalize_bill_df(agent.state.electricity_df)
    if df.empty:
        return []
    return df.where(pd.notna(df), None).to_dict(orient="records")


def _load_full_report_context(agent: StorageAgent, data: dict) -> dict:
    bills = _load_bills(agent) or _bill_payload(agent.state.electricity_df)
    capacity = _load_capacity(agent) or {}
    assumptions = capacity.setdefault("assumptions", {}) if isinstance(capacity, dict) else {}
    if isinstance(assumptions, dict):
        assumptions.setdefault("investment_unit_cost_yuan_per_kwh", DEFAULT_STORAGE_SYSTEM_COST_YUAN_PER_KWH)
        assumptions.setdefault("cell_cost_yuan_per_wh", DEFAULT_NA_CELL_COST_YUAN_PER_WH)
        assumptions.setdefault("cost_basis", DEFAULT_STORAGE_COST_BASIS)
    return {
        "revenue": data,
        "bills": bills,
        "bill_records": _bill_records_for_report(agent),
        "capacity": capacity,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }


def _safe_chart_text(text: str) -> str:
    return text.replace("·", "-")


def _generate_report_charts(ctx: dict, charts_dir: Path) -> dict[str, Path]:
    charts_dir.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib import font_manager
    except Exception as e:
        logger.exception("加载图表库失败: %s", e)
        return paths

    font_candidates = [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for font_path in font_candidates:
        if Path(font_path).exists():
            font_manager.fontManager.addfont(font_path)
            plt.rcParams["font.sans-serif"] = [font_manager.FontProperties(fname=font_path).get_name()]
            break
    plt.rcParams["axes.unicode_minus"] = False

    def savefig(name: str):
        path = charts_dir / f"{name}.png"
        plt.tight_layout()
        plt.savefig(path, dpi=180, bbox_inches="tight")
        plt.close()
        paths[name] = path

    bill_rows = ctx.get("bill_records") or []
    if bill_rows:
        months = [str(r.get("月份") or "") for r in bill_rows]
        plt.figure(figsize=(10.5, 4.2))
        plt.plot(months, [float(r.get("总电量(kWh)") or 0) / 10000 for r in bill_rows], marker="o", label="总电量(万kWh)")
        plt.plot(months, [float(r.get("总电费(元)") or 0) / 10000 for r in bill_rows], marker="s", label="总电费(万元)")
        plt.title("月度电量与电费")
        plt.xticks(rotation=35, ha="right")
        plt.grid(alpha=0.25)
        plt.legend()
        savefig("monthly_bill")

        plt.figure(figsize=(10.5, 4.2))
        plt.plot(months, [float(r.get("尖峰电价(元/kWh)") or 0) for r in bill_rows], marker="o", label="尖峰")
        plt.plot(months, [float(r.get("高峰电价(元/kWh)") or 0) for r in bill_rows], marker="o", label="高峰")
        plt.plot(months, [float(r.get("平段电价(元/kWh)") or 0) for r in bill_rows], marker="o", label="平段")
        plt.plot(months, [float(r.get("谷段电价(元/kWh)") or 0) for r in bill_rows], marker="o", label="谷段")
        plt.title("月度分时电价")
        plt.ylabel("元/kWh")
        plt.xticks(rotation=35, ha="right")
        plt.grid(alpha=0.25)
        plt.legend(ncol=4)
        savefig("tou_price")

        tou_labels = ["尖峰", "高峰", "平段", "谷段"]
        tou_values = [
            sum(float(r.get("尖峰电量(kWh)") or 0) for r in bill_rows),
            sum(float(r.get("高峰电量(kWh)") or 0) for r in bill_rows),
            sum(float(r.get("平段电量(kWh)") or 0) for r in bill_rows),
            sum(float(r.get("谷段电量(kWh)") or 0) for r in bill_rows),
        ]
        plt.figure(figsize=(7.2, 4.4))
        plt.pie(tou_values, labels=tou_labels, autopct="%1.1f%%", startangle=90)
        plt.title("全年分时用电结构")
        savefig("tou_energy_mix")

    capacity_rows = (ctx.get("capacity") or {}).get("results") or []
    if capacity_rows:
        top = sorted(capacity_rows, key=lambda r: float(r.get("battery_capacity_kwh") or 0))[:]
        labels = [str(r.get("name") or f"{r.get('battery_capacity_kwh')}kWh") for r in top]
        plt.figure(figsize=(10.8, 4.8))
        plt.bar(labels, [float(r.get("annual_revenue_yuan") or 0) / 10000 for r in top], label="年收益(万元)")
        plt.plot(labels, [float(r.get("payback_years") or 0) for r in top], color="#d14", marker="o", label="回收期(年)")
        plt.title("容量方案收益与回收期")
        plt.xticks(rotation=35, ha="right")
        plt.grid(axis="y", alpha=0.25)
        plt.legend()
        savefig("capacity_options")

    revenue = ctx.get("revenue") or {}
    cash_flow = revenue.get("cash_flow") or []
    if cash_flow:
        years = [int(r.get("year") or 0) for r in cash_flow]
        plt.figure(figsize=(10.5, 4.5))
        plt.bar(years, [float(r.get("investor_cash_flow_yuan") or 0) / 10000 for r in cash_flow], label="股东现金流(万元)")
        plt.plot(years, [float(r.get("cumulative_investor_yuan") or 0) / 10000 for r in cash_flow], color="#1a7", marker="o", label="累计股东现金流(万元)")
        plt.title("项目现金流曲线")
        plt.xlabel("年份")
        plt.grid(axis="y", alpha=0.25)
        plt.legend()
        savefig("cashflow_curve")

    sensitivity = revenue.get("sensitivity") or []
    if sensitivity:
        labels = [f"{r.get('factor')} {r.get('change_percent')}%" for r in sensitivity]
        plt.figure(figsize=(11, 5.2))
        plt.barh(labels, [float(r.get("payback_years") or 0) for r in sensitivity])
        plt.title("敏感性分析-回收期")
        plt.xlabel("回收期(年)")
        plt.grid(axis="x", alpha=0.25)
        savefig("sensitivity_payback")

    return paths


def _payback_matrix_table(matrix: dict) -> tuple[list[tuple[str, str]], list[dict]]:
    spreads = matrix.get("spreads") or []
    rows = matrix.get("rows") or []
    columns = [("unit_cost_yuan_per_wh", "系统+PCS单价/价差")]
    spread_keys = []
    for idx, spread in enumerate(spreads):
        key = f"spread_{idx}"
        spread_keys.append(key)
        columns.append((key, _md_value(spread, 2)))
    out_rows = []
    for row in rows:
        item = {"unit_cost_yuan_per_wh": row.get("unit_cost_yuan_per_wh")}
        for key, value in zip(spread_keys, row.get("values") or []):
            item[key] = value
        out_rows.append(item)
    return columns, out_rows


def _report_sections(ctx: dict, chart_paths: dict[str, Path] | None = None) -> list[dict]:
    chart_paths = chart_paths or {}
    data = ctx.get("revenue") or {}
    bills = ctx.get("bills") or {}
    bill_summary = bills.get("summary") or {}
    bill_records = ctx.get("bill_records") or []
    capacity = ctx.get("capacity") or {}
    capacity_best = capacity.get("best") or {}
    capacity_rows = capacity.get("results") or []
    params = data.get("params") or {}
    summary = data.get("summary") or {}
    loan = data.get("loan") or {}
    loan_impact = data.get("loan_impact") or {}
    project_name = params.get("project_name") or "储能收益测算"
    payback_columns, payback_rows = _payback_matrix_table(data.get("payback_matrix") or {})

    def chart(name: str, title: str) -> dict | None:
        path = chart_paths.get(name)
        if not path:
            return None
        return {"type": "image", "title": title, "path": path}

    sections = [
        {
            "title": "项目概览",
            "intro": f"{project_name}，报告生成时间：{ctx.get('generated_at') or '-'}，用户空间：{ctx.get('user_id') or ''}",
            "tables": [{
                "title": "核心结论",
                "columns": [("item", "指标"), ("value", "结果")],
                "rows": [
                    {"item": "推荐/当前规模", "value": f"{_report_num(summary.get('power_kw'), 0)} kW / {_report_num(summary.get('capacity_kwh'), 0)} kWh / {_report_num(summary.get('duration_hours'), 2)} h"},
                    {"item": "总投资", "value": _report_money(summary.get("total_investment_yuan"))},
                    {"item": "单位投资", "value": f"{_report_num(summary.get('unit_investment_yuan_per_wh'), 4)} 元/Wh"},
                    {"item": "成本口径", "value": params.get("cost_basis") or DEFAULT_STORAGE_COST_BASIS},
                    {"item": "首年总收入", "value": _report_money(summary.get("first_year_gross_revenue_yuan"))},
                    {"item": "峰谷套利毛利", "value": _report_money(summary.get("arbitrage_revenue_yuan"))},
                    {"item": "需量收益", "value": _report_money(summary.get("demand_revenue_yuan"))},
                    {"item": "静态回收期", "value": f"{_report_num(summary.get('static_payback_years'), 2)} 年"},
                    {"item": "项目 NPV", "value": _report_money(summary.get("project_npv_yuan"))},
                    {"item": "项目 IRR", "value": _report_percent(summary.get("project_irr_percent"))},
                    {"item": "股东 IRR", "value": _report_percent(summary.get("investor_irr_percent"))},
                ],
            }],
        },
        {
            "title": "账单与分时电价数据",
            "intro": "本节汇总上传账单解析结果，作为容量推荐和收益测算的负荷、电价基础。",
            "tables": [
                {
                    "title": "账单汇总",
                    "columns": [("item", "指标"), ("value", "结果")],
                    "rows": [
                        {"item": "账单月份", "value": f"{bill_summary.get('start_month') or '-'} 至 {bill_summary.get('end_month') or '-'}"},
                        {"item": "月数", "value": _report_num(bill_summary.get("month_count"), 0)},
                        {"item": "全年用电量", "value": f"{_report_num(bill_summary.get('total_kwh'), 2)} kWh"},
                        {"item": "全年电费", "value": _report_money(bill_summary.get("total_amount"))},
                        {"item": "平均电价", "value": f"{_report_num(bill_summary.get('avg_unit_price'), 4)} 元/kWh"},
                        {"item": "最大需量", "value": f"{_report_num(bill_summary.get('max_demand_kw'), 2)} kW"},
                        {"item": "需量电价", "value": f"{_report_num((bill_summary.get('prices') or {}).get('demand'), 2)} 元/kW·月"},
                    ],
                },
                {
                    "title": "分时电量与电价汇总",
                    "columns": [("period", "时段"), ("energy", "电量(kWh)"), ("price", "加权/平均电价(元/kWh)")],
                    "rows": [
                        {"period": "尖峰", "energy": (bill_summary.get("tou") or {}).get("peak"), "price": (bill_summary.get("prices") or {}).get("peak")},
                        {"period": "高峰", "energy": (bill_summary.get("tou") or {}).get("high"), "price": (bill_summary.get("prices") or {}).get("high")},
                        {"period": "平段", "energy": (bill_summary.get("tou") or {}).get("flat"), "price": (bill_summary.get("prices") or {}).get("flat")},
                        {"period": "谷段", "energy": (bill_summary.get("tou") or {}).get("valley"), "price": (bill_summary.get("prices") or {}).get("valley")},
                    ],
                },
                {
                    "title": "月度账单明细",
                    "columns": [
                        ("月份", "月份"), ("总电量(kWh)", "总电量(kWh)"), ("总电费(元)", "总电费(元)"),
                        ("尖峰电量(kWh)", "尖峰电量"), ("高峰电量(kWh)", "高峰电量"), ("平段电量(kWh)", "平段电量"), ("谷段电量(kWh)", "谷段电量"),
                        ("尖峰电价(元/kWh)", "尖峰电价"), ("高峰电价(元/kWh)", "高峰电价"), ("平段电价(元/kWh)", "平段电价"), ("谷段电价(元/kWh)", "谷段电价"),
                        ("最大需量(kW)", "最大需量(kW)"), ("需量电价(元/kW·月)", "需量电价"),
                    ],
                    "rows": bill_records,
                },
            ],
            "images": [i for i in [
                chart("monthly_bill", "月度电量与电费曲线"),
                chart("tou_price", "月度分时电价曲线"),
                chart("tou_energy_mix", "全年分时用电结构"),
            ] if i],
        },
        {
            "title": "储能最佳容量测算",
            "intro": capacity.get("scoring_basis") or "按容量候选方案对比年收益、单位容量收益、回收期、NPV 和 IRR，选择边际收益较优的容量点。",
            "tables": [
                {
                    "title": "推荐容量方案",
                    "columns": [("item", "指标"), ("value", "结果")],
                    "rows": [
                        {"item": "方案名称", "value": capacity_best.get("name") or "-"},
                        {"item": "成本口径", "value": (capacity.get("assumptions") or {}).get("cost_basis") or DEFAULT_STORAGE_COST_BASIS},
                        {"item": "钠离子电芯价格", "value": f"{_report_num((capacity.get('assumptions') or {}).get('cell_cost_yuan_per_wh'), 4)} 元/Wh"},
                        {"item": "完整系统投资单价", "value": f"{_report_num((capacity.get('assumptions') or {}).get('investment_unit_cost_yuan_per_kwh'), 2)} 元/kWh"},
                        {"item": "额定容量", "value": f"{_report_num(capacity_best.get('battery_capacity_kwh'), 2)} kWh"},
                        {"item": "装机功率", "value": f"{_report_num(capacity_best.get('inverter_power_kw'), 2)} kW"},
                        {"item": "储能时长", "value": f"{_report_num(capacity_best.get('duration_hours'), 2)} h"},
                        {"item": "年综合收益", "value": _report_money(capacity_best.get("annual_revenue_yuan"))},
                        {"item": "峰谷套利收益", "value": _report_money(capacity_best.get("arbitrage_revenue_yuan"))},
                        {"item": "需量收益", "value": _report_money(capacity_best.get("demand_revenue_yuan"))},
                        {"item": "收益/容量", "value": f"{_report_num(capacity_best.get('annual_revenue_per_kwh'), 2)} 元/kWh·年"},
                        {"item": "静态回收期", "value": f"{_report_num(capacity_best.get('payback_years'), 2)} 年"},
                        {"item": "选择理由", "value": capacity_best.get("selection_reason") or "-"},
                    ],
                },
                {
                    "title": "容量方案对比",
                    "columns": [
                        ("rank", "排名"), ("name", "配置"), ("battery_capacity_kwh", "容量(kWh)"), ("inverter_power_kw", "功率(kW)"),
                        ("duration_hours", "时长(h)"), ("annual_revenue_yuan", "年收益(元)"), ("annual_revenue_per_kwh", "收益/容量"),
                        ("payback_years", "回收期(年)"), ("npv_yuan", "NPV(元)"), ("irr_percent", "IRR(%)"),
                    ],
                    "rows": capacity_rows,
                },
                {
                    "title": "推荐方案月度套利测算",
                    "columns": [
                        ("month", "月份"), ("high_load_kwh", "高价负荷(kWh)"), ("discharge_kwh", "放电量(kWh)"),
                        ("charge_kwh", "充电量(kWh)"), ("discharge_value_yuan", "放电价值(元)"), ("charge_cost_yuan", "充电成本(元)"),
                        ("arbitrage_revenue_yuan", "套利收益(元)"), ("peak_price_yuan_per_kwh", "尖峰电价"), ("high_price_yuan_per_kwh", "高峰电价"),
                        ("valley_price_yuan_per_kwh", "谷电价"),
                    ],
                    "rows": capacity_best.get("monthly_estimate") or [],
                },
            ],
            "images": [i for i in [chart("capacity_options", "容量方案收益与回收期")] if i],
        },
        {
            "title": "收益测算参数与投资假设",
            "intro": "本节列示收益模型采用的电价、循环、成本、分成和贷款参数。",
            "tables": [
                {
                    "title": "关键输入参数",
                    "columns": [("item", "参数"), ("value", "取值")],
                    "rows": [
                        {"item": "装机功率", "value": f"{_report_num(params.get('power_kw'), 0)} kW"},
                        {"item": "额定容量", "value": f"{_report_num(params.get('capacity_kwh'), 0)} kWh"},
                        {"item": "钠离子电芯价格", "value": f"{_report_num(params.get('cell_cost_yuan_per_wh'), 4)} 元/Wh"},
                        {"item": "储能系统单价", "value": f"{_report_num(params.get('system_unit_cost_yuan_per_wh'), 4)} 元/Wh"},
                        {"item": "DOD", "value": _report_num(params.get("dod"), 4)},
                        {"item": "系统效率", "value": _report_num(params.get("system_efficiency"), 4)},
                        {"item": "年运行天数", "value": _report_num(params.get("annual_operating_days"), 0)},
                        {"item": "谷电充电电价", "value": f"{_report_num(params.get('valley_charge_price'), 4)} 元/kWh"},
                        {"item": "平电充电电价", "value": f"{_report_num(params.get('flat_charge_price'), 4)} 元/kWh"},
                        {"item": "放电电价", "value": f"{_report_num(params.get('discharge_price'), 4)} 元/kWh"},
                        {"item": "谷峰循环次数", "value": _report_num(params.get("valley_peak_cycles"), 3)},
                        {"item": "需量收益单价", "value": f"{_report_num(params.get('demand_revenue_per_kw_year'), 2)} 元/kW年"},
                        {"item": "客户分成", "value": "启用" if params.get("enable_customer_share") else "未启用"},
                        {"item": "贷款", "value": "启用" if params.get("enable_loan") else "未启用"},
                    ],
                },
                {
                    "title": "投资成本构成",
                    "columns": [("item", "成本项目"), ("amount_yuan", "金额(元)")],
                    "rows": data.get("cost_breakdown") or [],
                },
            ],
        },
        {
            "title": "投资收益与现金流分析",
            "intro": "本节展示年度项目现金流、客户收益、资方收益以及贷款影响。",
            "tables": [
                {
                    "title": "贷款影响",
                    "columns": [("item", "指标"), ("value", "结果")],
                    "rows": [
                        {"item": "贷款金额", "value": _report_money(loan.get("loan_amount_yuan"))},
                        {"item": "自有资金", "value": _report_money(loan.get("equity_amount_yuan"))},
                        {"item": "首年还本付息", "value": _report_money(loan.get("annual_debt_service_yuan"))},
                        {"item": "总利息", "value": _report_money(loan.get("total_interest_yuan"))},
                        {"item": "建议贷款比例", "value": loan.get("recommended_ratio_range") or "-"},
                        {"item": "贷款结论", "value": loan.get("conclusion") or "-"},
                        {"item": "有贷款股东 IRR", "value": _report_percent(loan_impact.get("loan_shareholder_irr_percent"))},
                        {"item": "无贷款股东 IRR", "value": _report_percent(loan_impact.get("no_loan_shareholder_irr_percent"))},
                    ],
                },
                {
                    "title": "有贷款 / 无贷款对比",
                    "columns": [("item", "指标"), ("value", "结果")],
                    "rows": [
                        {"item": "有贷款首年现金流", "value": _report_money(loan_impact.get("loan_first_year_cash_flow_yuan"))},
                        {"item": "无贷款首年现金流", "value": _report_money(loan_impact.get("no_loan_first_year_cash_flow_yuan"))},
                        {"item": "首年影响", "value": _report_money(loan_impact.get("first_year_delta_yuan"))},
                        {"item": "有贷款生命周期总收益", "value": _report_money(loan_impact.get("loan_lifetime_total_yuan"))},
                        {"item": "无贷款生命周期总收益", "value": _report_money(loan_impact.get("no_loan_lifetime_total_yuan"))},
                        {"item": "生命周期总收益影响", "value": _report_money(loan_impact.get("lifetime_delta_yuan"))},
                        {"item": "有贷款生命周期净现金流", "value": _report_money(loan_impact.get("loan_lifetime_net_cash_flow_yuan"))},
                        {"item": "无贷款生命周期净现金流", "value": _report_money(loan_impact.get("no_loan_lifetime_net_cash_flow_yuan"))},
                        {"item": "生命周期净现金流影响", "value": _report_money(loan_impact.get("net_delta_yuan"))},
                        {"item": "有贷款股东 IRR", "value": _report_percent(loan_impact.get("loan_shareholder_irr_percent"))},
                        {"item": "无贷款股东 IRR", "value": _report_percent(loan_impact.get("no_loan_shareholder_irr_percent"))},
                        {"item": "IRR 影响", "value": _report_percent(loan_impact.get("irr_delta_percent"))},
                    ],
                },
                {
                    "title": "项目现金流表",
                    "columns": [
                        ("year", "年份"), ("discharge_kwh", "放电量(kWh)"), ("charge_kwh", "充电量(kWh)"),
                        ("sale_revenue_yuan", "售电收入"), ("charge_cost_yuan", "充电成本"), ("arbitrage_margin_yuan", "套利毛利"),
                        ("demand_revenue_yuan", "需量收益"), ("gross_revenue_yuan", "总收入"), ("om_cost_yuan", "运维/保险"),
                        ("customer_share_yuan", "客户分成"), ("investor_ebitda_yuan", "资方EBITDA"), ("tax_yuan", "所得税"),
                        ("loan_interest_yuan", "利息"), ("loan_principal_yuan", "还本"), ("investor_cash_flow_yuan", "股东现金流"),
                        ("cumulative_investor_yuan", "累计股东现金流"),
                    ],
                    "rows": data.get("cash_flow") or [],
                },
                {
                    "title": "客户收益表",
                    "columns": [
                        ("year", "年份"), ("loan_status", "贷款状态"), ("share_ratio_percent", "分成比例(%)"),
                        ("distributable_yuan", "可分配收益"), ("customer_income_yuan", "客户年度收益"),
                        ("cumulative_customer_yuan", "累计客户收益"), ("loan_impact_yuan", "贷款影响"),
                    ],
                    "rows": data.get("customer_yearly") or [],
                },
                {
                    "title": "资方收益表",
                    "columns": [
                        ("year", "年份"), ("loan_status", "贷款状态"), ("loan_cost_yuan", "贷款成本"),
                        ("distributable_yuan", "可分配收益"), ("investor_income_yuan", "资方年度收益"),
                        ("cumulative_investor_income_yuan", "累计资方收益"), ("loan_impact_yuan", "贷款影响"),
                    ],
                    "rows": data.get("investor_yearly") or [],
                },
            ],
            "images": [i for i in [chart("cashflow_curve", "项目现金流曲线")] if i],
        },
        {
            "title": "敏感性分析",
            "intro": "本节用于观察电价、投资成本、循环次数、分成比例、贷款比例等关键变量变化对收益和回收期的影响。",
            "tables": [
                {
                    "title": "敏感性分析",
                    "columns": [("factor", "因素"), ("change_percent", "变化幅度(%)"), ("annual_net_yuan", "年净收益(元)"), ("payback_years", "回收期(年)")],
                    "rows": data.get("sensitivity") or [],
                },
                {
                    "title": "循环次数敏感性",
                    "columns": [
                        ("annual_cycles", "年循环次数"), ("discharge_kwh", "放电量(kWh)"), ("arbitrage_revenue_yuan", "套利毛利(元)"),
                        ("demand_revenue_yuan", "需量收益(元)"), ("annual_net_yuan", "年净收益(元)"), ("payback_years", "回收期(年)"),
                    ],
                    "rows": data.get("cycle_sensitivity") or [],
                },
                {
                    "title": "分成比例敏感性",
                    "columns": [("share_ratio_percent", "回本前客户分成(%)"), ("payback_years", "静态回收期(年)")],
                    "rows": data.get("share_sensitivity") or [],
                },
                {
                    "title": "静态回收期矩阵",
                    "columns": payback_columns,
                    "rows": payback_rows,
                },
            ],
            "images": [i for i in [chart("sensitivity_payback", "敏感性分析回收期曲线")] if i],
        },
        {
            "title": "测算口径",
            "bullets": [
                "峰谷套利收益 = 放电量 × 放电电价 - 充电量 × 充电电价。",
                "充电量 = 放电量 / 系统回路效率。",
                "容量测算会约束高价时段实际负荷承接能力，容量并非越大越优。",
                "需量收益按削峰能力、需量电价与兑现率估算；若没有 15 分钟负荷曲线，需量收益应保守使用。",
                "投资收益模型将客户分成、贷款、折旧、所得税、残值和电池更换成本纳入年度现金流。",
            ],
        },
    ]
    return sections


def _build_revenue_report_markdown(agent: StorageAgent, data: dict, report_dir: Path | None = None) -> str:
    ctx = _load_full_report_context(agent, data)
    ctx["user_id"] = agent.user_id
    chart_paths = _generate_report_charts(ctx, report_dir / "images") if report_dir else {}
    sections = _report_sections(ctx, chart_paths)
    params = data.get("params") or {}
    project_name = params.get("project_name") or "储能收益测算"
    parts = [
        f"# {project_name}详细测算报告",
        "",
        f"- 生成时间：{ctx.get('generated_at')}",
        f"- 用户空间：{agent.user_id}",
        f"- 模型版本：{data.get('model_version') or '-'}",
        f"- 模板来源：{data.get('model_source') or '-'}",
        "",
    ]
    for idx, section in enumerate(sections, start=1):
        parts.append(f"## {idx}. {section['title']}")
        if section.get("intro"):
            parts.extend(["", section["intro"], ""])
        for bullet in section.get("bullets") or []:
            parts.append(f"- {bullet}")
        if section.get("bullets"):
            parts.append("")
        for image in section.get("images") or []:
            rel = image["path"].relative_to(report_dir).as_posix() if report_dir else str(image["path"])
            parts.extend([f"![{image['title']}]({rel})", ""])
        for table in section.get("tables") or []:
            parts.extend([f"### {table['title']}", ""])
            parts.append(_md_table(table.get("rows") or [], table.get("columns") or []))
            parts.append("")
    return "\n".join(parts)


def _set_doc_table_style(table):
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Pt

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)


def _add_doc_table(document, rows: list[dict], columns: list[tuple[str, str]], *, limit: int | None = None):
    data = rows[:limit] if limit else rows
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, (_, label) in enumerate(columns):
        hdr[idx].text = str(label)
    if data:
        for row in data:
            cells = table.add_row().cells
            for idx, (key, _) in enumerate(columns):
                cells[idx].text = _md_value(row.get(key), 2)
    else:
        cells = table.add_row().cells
        cells[0].text = "暂无数据"
        for idx in range(1, len(cells)):
            cells[idx].text = ""
    _set_doc_table_style(table)
    return table


def _build_revenue_report_docx(agent: StorageAgent, data: dict, report_dir: Path, md_content: str | None = None) -> tuple[Path, str]:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.shared import Cm, Inches, Pt

    ctx = _load_full_report_context(agent, data)
    ctx["user_id"] = agent.user_id
    chart_paths = _generate_report_charts(ctx, report_dir / "images")
    if md_content is None:
        md_content = _build_revenue_report_markdown(agent, data, report_dir)
    sections = _report_sections(ctx, chart_paths)
    params = data.get("params") or {}
    project_name = params.get("project_name") or "储能收益测算"

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(9)

    document.add_heading(f"{project_name}详细测算报告", level=0)
    meta = document.add_paragraph()
    meta.add_run(f"生成时间：{ctx.get('generated_at')}    用户空间：{agent.user_id}").bold = True
    document.add_paragraph(f"模型版本：{data.get('model_version') or '-'}")
    document.add_paragraph(f"模板来源：{data.get('model_source') or '-'}")

    for idx, section_data in enumerate(sections, start=1):
        document.add_heading(f"{idx}. {section_data['title']}", level=1)
        if section_data.get("intro"):
            document.add_paragraph(str(section_data["intro"]))
        for bullet in section_data.get("bullets") or []:
            document.add_paragraph(str(bullet), style="List Bullet")
        for image in section_data.get("images") or []:
            document.add_paragraph(str(image["title"]))
            try:
                document.add_picture(str(image["path"]), width=Inches(8.8))
            except Exception:
                logger.exception("插入报告图片失败: %s", image.get("path"))
        for table in section_data.get("tables") or []:
            document.add_heading(str(table["title"]), level=2)
            _add_doc_table(document, table.get("rows") or [], table.get("columns") or [])

    filename = f"storage_full_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    path = report_dir / filename
    document.save(path)
    return path, md_content


def _build_revenue_report_package(agent: StorageAgent, data: dict) -> dict:
    reports_dir = DATA_ROOT / "output" / "reports" / agent.user_id
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_dir = reports_dir / f"storage_full_report_{stamp}"
    report_dir.mkdir(parents=True, exist_ok=True)

    md_content = _build_revenue_report_markdown(agent, data, report_dir)
    md_path = report_dir / f"storage_full_report_{stamp}.md"
    md_path.write_text(md_content, encoding="utf-8")
    docx_path, _ = _build_revenue_report_docx(agent, data, report_dir, md_content)

    zip_path = reports_dir / f"storage_full_report_{stamp}.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.write(md_path, md_path.relative_to(report_dir.parent))
        zf.write(docx_path, docx_path.relative_to(report_dir.parent))
        images_dir = report_dir / "images"
        if images_dir.exists():
            for image in images_dir.glob("*.png"):
                zf.write(image, image.relative_to(report_dir.parent))
    return {
        "ok": True,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "report_dir": str(report_dir),
        "zip_path": str(zip_path),
        "docx_path": str(docx_path),
        "md_path": str(md_path),
        "zip_url": f"/api/reports/download?user_id={agent.user_id}&path={zip_path.relative_to(reports_dir).as_posix()}",
        "docx_url": f"/api/reports/download?user_id={agent.user_id}&path={docx_path.relative_to(reports_dir).as_posix()}",
        "md_url": f"/api/reports/download?user_id={agent.user_id}&path={md_path.relative_to(reports_dir).as_posix()}",
    }


# ======================================================================
# FastAPI App
# ======================================================================
def create_app(manager: AgentManager) -> FastAPI:
    app = FastAPI(title="储能 AGENT", version="2.0")

    app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_methods=["*"],
        allow_headers=["*"],
    )

    if STATIC_DIR.exists():
        app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")

    # --------------------------------------------------------------
    # 主页
    # --------------------------------------------------------------
    @app.get("/")
    async def index():
        idx = STATIC_DIR / "index.html"
        if idx.exists():
            return FileResponse(str(idx))
        return JSONResponse({"error": "static/index.html not found"}, status_code=500)

    @app.get("/bills")
    @app.get("/bills/")
    async def bills_page():
        page = STATIC_DIR / "bills.html"
        if page.exists():
            return FileResponse(str(page))
        return JSONResponse({"error": "static/bills.html not found"}, status_code=500)

    @app.get("/capacity")
    @app.get("/capacity/")
    async def capacity_page():
        page = STATIC_DIR / "capacity.html"
        if page.exists():
            return FileResponse(str(page))
        return JSONResponse({"error": "static/capacity.html not found"}, status_code=500)

    @app.get("/revenue")
    @app.get("/revenue/")
    async def revenue_page():
        page = STATIC_DIR / "revenue.html"
        if page.exists():
            return FileResponse(str(page))
        return JSONResponse({"error": "static/revenue.html not found"}, status_code=500)

    # --------------------------------------------------------------
    # 用户 / 状态 / 记忆 / 知识库
    # --------------------------------------------------------------
    @app.get("/api/me")
    async def get_me(x_auth_user: str | None = Header(default=None)):
        user_id = _effective_user_id(None, x_auth_user)
        return {"user_id": user_id, "auth_user": x_auth_user or "", "users": manager.list_users()}

    @app.get("/api/users")
    async def get_users(x_auth_user: str | None = Header(default=None)):
        current_user = _effective_user_id(None, x_auth_user)
        users = manager.list_users()
        if current_user not in users:
            users.append(current_user)
            users = sorted(set(users), key=str.lower)
        return {"users": users, "current_user": current_user, "auth_user": x_auth_user or ""}

    @app.post("/api/users")
    async def create_user(req: dict):
        uid = (req.get("user_id") or "").strip()
        if not uid:
            raise HTTPException(400, "user_id required")
        manager.get(uid)
        return {"ok": True, "user_id": safe_user_id(uid), "users": manager.list_users()}

    @app.get("/api/state")
    async def get_state(user_id: str | None = None, x_auth_user: str | None = Header(default=None)):
        agent = await asyncio.to_thread(manager.get, _effective_user_id(user_id, x_auth_user))
        return _state_summary(agent)

    @app.get("/api/memory")
    async def get_memory(user_id: str | None = None, x_auth_user: str | None = Header(default=None)):
        agent = await asyncio.to_thread(manager.get, _effective_user_id(user_id, x_auth_user))
        return _memory_summary(agent)

    @app.get("/api/kb")
    async def get_kb(user_id: str | None = None, x_auth_user: str | None = Header(default=None)):
        agent = await asyncio.to_thread(manager.get, _effective_user_id(user_id, x_auth_user))
        return _kb_summary(agent)

    @app.get("/api/history")
    async def get_history(user_id: str | None = None, limit: int = 100, x_auth_user: str | None = Header(default=None)):
        """返回当前会话的 user/assistant 消息历史，用于浏览器刷新/切用户后恢复显示。

        - 跳过 system prompt 与 tool 消息（这些不直接展示给用户）
        - 同时附带每个 assistant 消息已知的工具调用名（仅展示用，不还原工具结果详情）
        """
        agent = await asyncio.to_thread(manager.get, _effective_user_id(user_id, x_auth_user))
        msgs = list(getattr(agent, "messages", []) or [])
        out: list[dict] = []
        # 跳过首条 system
        for m in msgs:
            if not isinstance(m, dict):
                # 可能是 OpenAI SDK 对象，转 dict
                try:
                    m = m.model_dump()  # pydantic v2
                except Exception:
                    try:
                        m = m.dict()
                    except Exception:
                        m = {"role": getattr(m, "role", "?"),
                             "content": getattr(m, "content", "") or ""}
            role = m.get("role")
            if role == "system":
                continue
            if role == "tool":
                continue
            content = m.get("content") or ""
            tool_calls = m.get("tool_calls") or []
            tool_names = []
            try:
                for tc in tool_calls:
                    if isinstance(tc, dict):
                        tool_names.append((tc.get("function") or {}).get("name") or tc.get("name"))
                    else:
                        fn = getattr(tc, "function", None)
                        tool_names.append(getattr(fn, "name", None) if fn else None)
            except Exception:
                pass
            tool_names = [t for t in tool_names if t]
            # 不展示空 assistant + 无 tool_calls 的占位项
            if role == "assistant" and not content and not tool_names:
                continue
            out.append({
                "role": role,
                "content": content,
                "tool_calls": tool_names,
            })
        if limit and len(out) > limit:
            out = out[-limit:]
        return {"messages": out, "total": len(out)}

    # --------------------------------------------------------------
    # 聊天（SSE 流式）
    # --------------------------------------------------------------
    @app.post("/api/chat")
    async def chat(req: dict, x_auth_user: str | None = Header(default=None)):
        user_id = _effective_user_id(req.get("user_id"), x_auth_user)
        message = (req.get("message") or "").strip()
        if not message:
            raise HTTPException(400, "empty message")

        agent = await asyncio.to_thread(manager.get, user_id)
        loop = asyncio.get_event_loop()

        async def event_stream() -> AsyncGenerator[str, None]:
            queue: asyncio.Queue = asyncio.Queue()
            SENTINEL = object()

            def producer():
                try:
                    for ev in agent.chat_stream(message):
                        # 过滤内部魔法事件
                        et = ev.get("type") if isinstance(ev, dict) else None
                        if et and et.startswith("_"):
                            continue
                        loop.call_soon_threadsafe(queue.put_nowait, ev)
                except Exception as e:
                    logger.exception("chat_stream 异常")
                    loop.call_soon_threadsafe(
                        queue.put_nowait, {"type": "error", "message": str(e)}
                    )
                finally:
                    try:
                        _persist_runtime_state(agent)
                    except Exception:
                        logger.exception("保存 Agent 运行状态失败")
                    loop.call_soon_threadsafe(queue.put_nowait, SENTINEL)

            t = threading.Thread(target=producer, daemon=True)
            t.start()

            try:
                while True:
                    ev = await queue.get()
                    if ev is SENTINEL:
                        break
                    payload = json.dumps(ev, ensure_ascii=False, default=str)
                    yield f"data: {payload}\n\n"
                yield f"data: {json.dumps({'type': 'done'})}\n\n"
            except asyncio.CancelledError:
                logger.info("SSE 客户端断开")
                raise

        return StreamingResponse(
            event_stream(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",  # 禁用 nginx buffering
            },
        )

    @app.post("/api/clear")
    async def clear_chat(req: dict, x_auth_user: str | None = Header(default=None)):
        user_id = _effective_user_id(req.get("user_id"), x_auth_user)
        agent = await asyncio.to_thread(manager.get, user_id)
        # 仅清当前会话上下文（保留长期记忆）
        try:
            agent.messages = agent.messages[:1]  # 保留 system prompt
        except Exception:
            pass
        return {"ok": True}

    @app.post("/api/reset")
    async def reset_all(req: dict, x_auth_user: str | None = Header(default=None)):
        user_id = _effective_user_id(req.get("user_id"), x_auth_user)
        agent = await asyncio.to_thread(manager.get, user_id)
        await asyncio.to_thread(agent.reset_all)
        await asyncio.to_thread(_clear_persisted_state, agent)
        return {"ok": True, "state": _state_summary(agent)}

    # --------------------------------------------------------------
    # 电费账单解析 / 表格视图
    # --------------------------------------------------------------
    @app.get("/api/bills")
    async def get_bills(user_id: str | None = None, x_auth_user: str | None = Header(default=None)):
        agent = await asyncio.to_thread(manager.get, _effective_user_id(user_id, x_auth_user))
        cached = await asyncio.to_thread(_load_bills, agent)
        if cached:
            return cached
        return _bill_payload(agent.state.electricity_df)

    @app.post("/api/bills/parse")
    async def parse_bills(req: dict, x_auth_user: str | None = Header(default=None)):
        user_id = _effective_user_id(req.get("user_id"), x_auth_user)
        mode = (req.get("mode") or "auto").strip().lower()
        if mode not in {"auto", "llm", "rules"}:
            raise HTTPException(400, "mode 只能是 auto / llm / rules")

        names = req.get("files") or []
        if isinstance(names, str):
            names = [names]
        if not names:
            names = _list_bill_files()
        if not names:
            raise HTTPException(400, "input/ 目录没有可解析的账单文件")

        paths = [_resolve_input_file(str(name)) for name in names]
        agent = await asyncio.to_thread(manager.get, user_id)
        df, parser_used = await asyncio.to_thread(_parse_bill_paths_sync, agent, paths, mode)
        if df.empty:
            raise HTTPException(422, "未能从选中文件中提取到有效电费数据")
        return _bill_payload(
            df,
            files=[p.name for p in paths],
            parser=parser_used,
            msg=f"已解析 {len(paths)} 个文件，提取到 {len(df)} 条账单记录",
        )

    @app.post("/api/bills/parse/stream")
    async def parse_bills_stream(req: dict, x_auth_user: str | None = Header(default=None)):
        user_id = _effective_user_id(req.get("user_id"), x_auth_user)
        mode = (req.get("mode") or "auto").strip().lower()
        if mode not in {"auto", "llm", "rules"}:
            raise HTTPException(400, "mode 只能是 auto / llm / rules")

        names = req.get("files") or []
        if isinstance(names, str):
            names = [names]
        if not names:
            names = _list_bill_files()
        if not names:
            raise HTTPException(400, "input/ 目录没有可解析的账单文件")

        paths = [_resolve_input_file(str(name)) for name in names]
        agent = await asyncio.to_thread(manager.get, user_id)

        async def event_stream() -> AsyncGenerator[str, None]:
            total = len(paths)
            frames: list[pd.DataFrame] = []

            async def emit(payload: dict):
                data = json.dumps(payload, ensure_ascii=False, default=str)
                return f"data: {data}\n\n"

            yield await emit({
                "type": "start",
                "total": total,
                "files": [p.name for p in paths],
            })

            for idx, path in enumerate(paths, start=1):
                yield await emit({
                    "type": "file_start",
                    "index": idx,
                    "total": total,
                    "file": path.name,
                    "progress": round((idx - 1) / total * 100, 1),
                })
                try:
                    df, parser_used = await asyncio.to_thread(
                        _parse_single_bill_path_sync, agent, path, mode
                    )
                    if df.empty:
                        yield await emit({
                            "type": "file_error",
                            "index": idx,
                            "total": total,
                            "file": path.name,
                            "message": "未提取到有效账单数据",
                            "progress": round(idx / total * 100, 1),
                        })
                        continue

                    frames.append(df)
                    combined = _set_bill_state(
                        agent,
                        pd.concat(frames, ignore_index=True),
                        files=[p.name for p in paths[:idx]],
                        parser=parser_used,
                        msg=f"已解析 {idx}/{total} 个文件",
                        merge=True,
                    )
                    payload = _bill_payload(
                        combined,
                        files=[p.name for p in paths[:idx]],
                        parser=parser_used,
                        msg=f"已解析 {idx}/{total} 个文件",
                    )
                    yield await emit({
                        "type": "file_done",
                        "index": idx,
                        "total": total,
                        "file": path.name,
                        "parser": parser_used,
                        "rows_added": int(len(df)),
                        "progress": round(idx / total * 100, 1),
                        "bill": _bill_payload(df, files=[path.name], parser=parser_used),
                        "payload": payload,
                    })
                except Exception as e:
                    logger.exception("账单解析失败: %s", path.name)
                    yield await emit({
                        "type": "file_error",
                        "index": idx,
                        "total": total,
                        "file": path.name,
                        "message": str(e),
                        "progress": round(idx / total * 100, 1),
                    })

            if frames:
                combined = _set_bill_state(
                    agent,
                    pd.concat(frames, ignore_index=True),
                    files=[p.name for p in paths],
                    msg=f"已完成 {len(frames)}/{total} 个文件",
                    merge=True,
                )
                final_payload = _bill_payload(
                    combined,
                    files=[p.name for p in paths],
                    msg=f"已完成 {len(frames)}/{total} 个文件",
                )
                yield await emit({
                    "type": "done",
                    "total": total,
                    "success": len(frames),
                    "failed": total - len(frames),
                    "progress": 100,
                    "payload": final_payload,
                })
            else:
                yield await emit({
                    "type": "done",
                    "total": total,
                    "success": 0,
                    "failed": total,
                    "progress": 100,
                    "payload": _bill_payload(None, files=[p.name for p in paths]),
                })

        return StreamingResponse(
            event_stream(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
            },
        )

    @app.post("/api/storage/docs/analyze")
    async def analyze_storage_docs(req: dict, x_auth_user: str | None = Header(default=None)):
        user_id = _effective_user_id(req.get("user_id"), x_auth_user)
        names = req.get("files") or req.get("file_paths") or []
        if isinstance(names, str):
            names = [names]
        if not names:
            raise HTTPException(400, "files 不能为空")
        topic = (req.get("topic") or "分时电价、需量电价、容量电费、储能套利、削峰填谷、补贴政策、并网要求").strip()
        index_to_kb = _bool_param(req.get("index_to_kb", True))
        use_llm = _bool_param(req.get("use_llm", False))
        paths = [_resolve_input_file(str(name)) for name in names]
        agent = await asyncio.to_thread(manager.get, user_id)
        return await asyncio.to_thread(_analyze_storage_docs_sync, agent, paths, topic, index_to_kb, use_llm)

    @app.post("/api/storage/capacity-analysis")
    async def capacity_analysis(req: dict, x_auth_user: str | None = Header(default=None)):
        user_id = _effective_user_id(req.get("user_id"), x_auth_user)
        agent = await asyncio.to_thread(manager.get, user_id)
        return await asyncio.to_thread(_capacity_analysis_sync, agent, req)

    @app.get("/api/storage/capacity-analysis")
    async def get_capacity_analysis(user_id: str | None = None, x_auth_user: str | None = Header(default=None)):
        agent = await asyncio.to_thread(manager.get, _effective_user_id(user_id, x_auth_user))
        data = await asyncio.to_thread(_load_capacity, agent)
        if data:
            return data
        rows = getattr(agent.state, "capacity_analysis", None) or []
        if rows:
            return {
                "ok": True,
                "msg": "已加载当前容量分析结果",
                "candidate_count": len(rows),
                "positive_count": sum(1 for row in rows if (row.get("annual_revenue_yuan") or 0) > 0),
                "best": next((row for row in rows if row.get("is_best")), rows[0]),
                "results": rows,
            }
        return {"ok": True, "msg": "暂无容量分析结果", "candidate_count": 0, "positive_count": 0, "results": []}

    @app.get("/api/revenue/model")
    async def get_revenue_model(user_id: str | None = None, x_auth_user: str | None = Header(default=None)):
        user_id = _effective_user_id(user_id, x_auth_user)
        agent = await asyncio.to_thread(manager.get, user_id)
        data = await asyncio.to_thread(_load_revenue, agent)
        capacity = await asyncio.to_thread(_load_capacity, agent)
        if (
            data
            and data.get("model_version") == REVENUE_MODEL_VERSION
            and (data.get("params_source") == "manual" or not capacity or _revenue_matches_capacity(data, capacity))
        ):
            return data
        return await asyncio.to_thread(_refresh_revenue_for_current_capacity, agent, user_id)

    @app.post("/api/revenue/model")
    async def post_revenue_model(req: dict, x_auth_user: str | None = Header(default=None)):
        user_id = _effective_user_id(req.get("user_id"), x_auth_user)
        agent = await asyncio.to_thread(manager.get, user_id)
        return await asyncio.to_thread(_compute_revenue_model, agent, req)

    @app.get("/api/revenue/params/history")
    async def get_revenue_param_history(user_id: str | None = None, x_auth_user: str | None = Header(default=None)):
        user_id = _effective_user_id(user_id, x_auth_user)
        agent = await asyncio.to_thread(manager.get, user_id)
        return await asyncio.to_thread(_load_revenue_param_history, agent)

    @app.get("/api/revenue/report")
    async def get_revenue_report(user_id: str | None = None, x_auth_user: str | None = Header(default=None)):
        user_id = _effective_user_id(user_id, x_auth_user)
        agent = await asyncio.to_thread(manager.get, user_id)
        data = await asyncio.to_thread(_load_revenue, agent)
        if not data:
            data = await asyncio.to_thread(_refresh_revenue_for_current_capacity, agent, user_id)
        package = await asyncio.to_thread(_build_revenue_report_package, agent, data)
        path = Path(package["zip_path"])
        return FileResponse(
            str(path),
            media_type="application/zip",
            filename=path.name,
        )

    @app.post("/api/revenue/report")
    async def post_revenue_report(req: dict, x_auth_user: str | None = Header(default=None)):
        user_id = _effective_user_id(req.get("user_id"), x_auth_user)
        agent = await asyncio.to_thread(manager.get, user_id)
        data = await asyncio.to_thread(_load_revenue, agent)
        if not data:
            data = await asyncio.to_thread(_refresh_revenue_for_current_capacity, agent, user_id)
        return await asyncio.to_thread(_build_revenue_report_package, agent, data)

    @app.get("/api/reports/download")
    async def download_report_file(user_id: str | None = None, path: str = "", x_auth_user: str | None = Header(default=None)):
        user_id = _effective_user_id(user_id, x_auth_user)
        raw = str(path or "").strip()
        if not raw:
            raise HTTPException(400, "path 不能为空")
        rel = Path(raw)
        if rel.is_absolute() or ".." in rel.parts:
            raise HTTPException(400, "非法报告路径")
        reports_dir = (DATA_ROOT / "output" / "reports" / user_id).resolve()
        target = (reports_dir / rel).resolve()
        if reports_dir not in target.parents and target != reports_dir:
            raise HTTPException(400, "非法报告路径")
        if not target.exists() or not target.is_file():
            raise HTTPException(404, "报告文件不存在")
        if target.suffix.lower() not in REPORT_FILE_EXTS:
            raise HTTPException(400, "不支持下载的报告文件类型")
        media_types = {
            ".zip": "application/zip",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".md": "text/markdown; charset=utf-8",
            ".png": "image/png",
        }
        return FileResponse(str(target), media_type=media_types.get(target.suffix.lower()), filename=target.name)

    # --------------------------------------------------------------
    # 文件上传（input/）—— 多文件 / 文件夹
    # --------------------------------------------------------------
    @app.post("/api/upload")
    async def upload(
        files: list[UploadFile] = File(...),
        user_id: str | None = Form(None),
        from_folder: str = Form("0"),
        x_auth_user: str | None = Header(default=None),
    ):
        user_id = _effective_user_id(user_id, x_auth_user)
        from_folder_b = from_folder in ("1", "true", "True")
        copied: list[str] = []
        skipped: list[str] = []
        errors: list[str] = []
        # 文件夹模式：按白名单过滤；单文件模式：宽松（如果用户选了非白名单也允许）
        allowed = SUPPORTED_UPLOAD_EXTS if from_folder_b else None
        for f in files:
            dst, msg = _save_upload(f, allowed)
            if dst is None:
                if msg.startswith("不支持的扩展名"):
                    skipped.append(Path(f.filename or "").name)
                else:
                    errors.append(f"{f.filename}: {msg}")
                continue
            copied.append(dst.name)
        return {
            "ok": True,
            "copied": copied,
            "skipped": skipped,
            "errors": errors,
            "total": len(copied),
        }

    @app.delete("/api/input/{name}")
    async def delete_input_file(name: str):
        # 只允许删除 INPUT_DIR 下的文件
        target = (INPUT_DIR / name).resolve()
        if INPUT_DIR.resolve() not in target.parents and target != INPUT_DIR.resolve() / name:
            raise HTTPException(400, "非法路径")
        if not target.exists() or not target.is_file():
            raise HTTPException(404, "文件不存在")
        try:
            target.unlink()
            return {"ok": True}
        except Exception as e:
            raise HTTPException(500, str(e))

    # --------------------------------------------------------------
    # 知识库
    # --------------------------------------------------------------
    @app.post("/api/kb/index")
    async def kb_index(
        files: list[UploadFile] = File(...),
        user_id: str | None = Form(None),
        from_folder: str = Form("0"),
        x_auth_user: str | None = Header(default=None),
    ):
        from_folder_b = from_folder in ("1", "true", "True")
        user_id = _effective_user_id(user_id, x_auth_user)
        agent = await asyncio.to_thread(manager.get, user_id)
        kb = agent.state.kb
        if kb is None or not getattr(kb, "is_ready", False):
            raise HTTPException(400, "知识库未启用")

        results: list[dict] = []
        for f in files:
            name = Path(f.filename or "").name
            ext = Path(name).suffix.lower()
            if from_folder_b and ext not in SUPPORTED_KB_EXTS:
                results.append({"name": name, "ok": False, "msg": "扩展名不支持，已跳过"})
                continue
            # 写到临时目录再 index
            try:
                with tempfile.NamedTemporaryFile(
                    delete=False, suffix=ext or ".bin"
                ) as tmp:
                    shutil.copyfileobj(f.file, tmp)
                    tmp_path = Path(tmp.name)
            except Exception as e:
                results.append({"name": name, "ok": False, "msg": f"写入失败: {e}"})
                continue
            try:
                # 改名为原名以便 KB 内 source 字段是真实文件名
                final_path = tmp_path.parent / name
                try:
                    if final_path.exists():
                        final_path.unlink()
                    tmp_path.rename(final_path)
                except Exception:
                    final_path = tmp_path  # 回退
                n = await asyncio.to_thread(kb.index_file, str(final_path))
                results.append({"name": name, "ok": True, "chunks": int(n)})
            except Exception as e:
                results.append({"name": name, "ok": False, "msg": str(e)})
            finally:
                try:
                    if tmp_path.exists():
                        tmp_path.unlink()
                except Exception:
                    pass
                try:
                    if 'final_path' in locals() and final_path.exists() and final_path != tmp_path:
                        final_path.unlink()
                except Exception:
                    pass
        return {"results": results, "kb": _kb_summary(agent)}

    @app.post("/api/kb/search")
    async def kb_search(req: dict, x_auth_user: str | None = Header(default=None)):
        user_id = _effective_user_id(req.get("user_id"), x_auth_user)
        query = (req.get("query") or "").strip()
        k = int(req.get("k") or 5)
        if not query:
            raise HTTPException(400, "empty query")
        agent = await asyncio.to_thread(manager.get, user_id)
        kb = agent.state.kb
        if kb is None or not getattr(kb, "is_ready", False):
            raise HTTPException(400, "KB 未启用")
        try:
            hits = await asyncio.to_thread(kb.search, query, k)
        except TypeError:
            hits = await asyncio.to_thread(lambda: kb.search(query, k=k))
        # hits 元素可能是 dict 或自定义类
        norm = []
        for h in hits or []:
            if isinstance(h, dict):
                norm.append(h)
            else:
                norm.append({
                    "score": getattr(h, "score", None),
                    "text": getattr(h, "text", None) or getattr(h, "content", None),
                    "source": getattr(h, "source", None),
                })
        return {"hits": norm}

    @app.delete("/api/kb/{source}")
    async def kb_remove(source: str, user_id: str | None = None, x_auth_user: str | None = Header(default=None)):
        agent = await asyncio.to_thread(manager.get, _effective_user_id(user_id, x_auth_user))
        kb = agent.state.kb
        if kb is None or not getattr(kb, "is_ready", False):
            raise HTTPException(400, "KB 未启用")
        try:
            await asyncio.to_thread(kb.remove_document, source)
            return {"ok": True, "kb": _kb_summary(agent)}
        except Exception as e:
            raise HTTPException(500, str(e))

    return app


# ======================================================================
# 启动入口（被 main.py 调用）
# ======================================================================
def launch(config: AgentConfig, host: str = "127.0.0.1", port: int = 7860,
           default_user: str = "main", **kwargs):
    """启动 FastAPI Web UI。"""
    try:
        import uvicorn
    except ImportError:
        raise ImportError(
            "FastAPI Web UI 依赖 uvicorn，未安装。请运行：\n"
            "  pip install fastapi uvicorn[standard] python-multipart"
        )

    manager = AgentManager(config=config, **kwargs)
    # 预热默认用户
    try:
        manager.get(default_user)
    except Exception as e:
        logger.warning("预热默认用户失败: %s", e)

    app = create_app(manager)

    print()
    print("=" * 60)
    print("🚀 储能 AGENT Web 启动中（FastAPI / 异步无锁）")
    print(f"   访问: http://{host}:{port}")
    print(f"   默认用户: {default_user}")
    print("   按 Ctrl+C 停止")
    print("=" * 60)
    print()

    uvicorn.run(app, host=host, port=port, log_level="info", access_log=False)
